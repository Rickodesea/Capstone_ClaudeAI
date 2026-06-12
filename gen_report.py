"""
gen_report.py  --  Optimization layer report for DAMO 699 Capstone
Run:  python gen_report.py
Output: optimization_layer_report.docx
"""

from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ROOT    = Path(__file__).resolve().parent
PLOTS   = ROOT / "Cluster_Optimization_Models" / "Pipeline" / "timing_data" / "plots"
SCREENS = ROOT / "Cluster_Optimization_Models" / "Simulation" / "screenshots"
OUT     = ROOT / "optimization_layer_report.docx"


# ── Helpers ────────────────────────────────────────────────────────────────────

def _heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def _body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def _bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")

def _equation(doc, text):
    """
    Render an equation line using Cambria Math font, centered, slightly indented.
    python-docx does not expose the full OMML equation editor, so we use a
    formatted paragraph — sufficient for report quality output.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)
    return p

def _image(doc, path, width=5.5, caption=None):
    path = Path(path)
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Figure not found: {path.name}]")
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.italic = True
        run.font.size = Pt(9)

def _table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return t

def _page_break(doc):
    doc.add_page_break()


# ── Build ──────────────────────────────────────────────────────────────────────

doc = Document()

# ── Title ──────────────────────────────────────────────────────────────────────
title = doc.add_heading(
    "A Unified Framework for Multi-Tenant Cluster Management\n"
    "with Overcommit and Guaranteed Service Level Agreement",
    level=0
)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("Optimization Layer")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True

doc.add_paragraph("DAMO 699 — Capstone Project").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")
_page_break(doc)


# ── 1. Introduction ────────────────────────────────────────────────────────────
_heading(doc, "1. Introduction")

p = doc.add_paragraph()
p.add_run("Abstract.  ").bold = True
p.add_run(
    "This project develops a two-layer framework for managing multi-tenant cluster workloads. "
    "The prediction layer produces per-job memory and CPU predictions and per-tenant aggregate "
    "demand forecasts. The optimization layer consumes those predictions to solve two coupled "
    "problems: a periodic plan-ahead that partitions cluster nodes among tenants, and a "
    "real-time scheduler that assigns individual jobs to nodes each epoch. "
    "Both models are formulated as exact integer programs. Because single-shot integer programs "
    "become intractable at production scale, an iterative decomposition strategy is applied, "
    "reducing solve time from hours to seconds while maintaining 100 percent job placement rate."
)
p.paragraph_format.space_after = Pt(8)

_body(doc,
    "The framework connects two layers through a defined data contract. The prediction layer "
    "provides pred_mem_mb and pred_cpu_p95 per job for the real-time model, and u[i,h] "
    "and sigma2[i,h] per tenant per period for the plan-ahead model. The optimization layer "
    "uses these values directly in its objective and constraints. No retraining or coupling "
    "between layers is needed."
)

_body(doc,
    "The project uses a subset of the Google Cluster Usage Traces v3 dataset, covering nine "
    "tenants across a 24-hour planning horizon with six-hour planning slots. A Borg "
    "configuration module (borg_configuration.py) stores dataset-derived constants and can be "
    "loaded into the simulation dashboard to replace synthetic defaults with real dataset "
    "parameters."
)

_page_break(doc)


# ── 2. Analytical Methods ──────────────────────────────────────────────────────
_heading(doc, "2. Analytical Methods")

_heading(doc, "2.1  Real-Time Scheduler (MILP)", level=2)

_body(doc,
    "At each scheduling epoch the real-time model assigns pending jobs to cluster nodes. "
    "The binary decision variable x_{jn} equals 1 if job j is placed on node n."
)

_body(doc, "Objective — maximize weighted memory placement:")
_equation(doc, "max Z = Σ_j Σ_n  ω_t(j) * P̂_j * u_n * σ_n * x_{jn}")

_body(doc, "Constraints:")
_bullet(doc, "C1:  Σ_n x_{jn}  ≤  1          each job placed on at most one node")
_bullet(doc, "C2:  Σ_j P̂_jᵐᵉᵐ * x_{jn}  ≤  M_nᵉᶠᶠ     predicted memory fits within effective remaining capacity")
_bullet(doc, "C3:  x_{jn} ∈ {0,1}")
_bullet(doc, "C4:  x_{jn} = 0  if  P̂_jᶜᵖᵘ > C_n      CPU demand enforced as variable upper bound")

_body(doc,
    "The objective weight ω_t boosts priority for tenants whose rolling average wait "
    "exceeds the cluster mean, providing fairness as a side effect of weighted maximization "
    "(Shi and Yu, 2024). Effective capacity on each node shrinks when the node has recent "
    "SLA violations, preventing further overcommit on stressed nodes."
)
_body(doc,
    "Gurobi is the default backend. It consistently outperforms open-source alternatives "
    "at all problem sizes tested. At J=256 jobs and N=64 nodes, Gurobi reaches optimality in "
    "under 30 seconds where CBC requires over 120 seconds. The backend is selectable via "
    "configuration: CBC, SCIP, and HiGHS are also supported when a Gurobi license is "
    "unavailable."
)

_heading(doc, "2.2  Plan-Ahead Allocator (MISOCP)", level=2)

_body(doc,
    "The plan-ahead model runs periodically and partitions cluster nodes among tenants across "
    "a planning horizon H. Exclusive tenants receive dedicated machines per period; "
    "shared tenants are grouped together and share a pool of machines per period."
)

_body(doc,
    "Tenant grouping is the core of the plan-ahead step. Tenants are first split into exclusive "
    "and shared classes. Shared tenants are then partitioned by demand into heavy and light: a "
    "tenant is heavy if its average demand is at or above the median across shared tenants, and "
    "light otherwise. The objective's mix bonus (λ₂ * mix_total) rewards co-locating a heavy "
    "tenant and a light tenant on the same machine, so the model actively pairs high-load "
    "tenants with low-load tenants rather than packing similar loads together. This smooths "
    "aggregate demand per machine and leaves more headroom for the Cantelli safety buffer. "
    "Once tenants are grouped, the model assigns a set of machines to each group for each "
    "period; the y[i,n,h] variables are the hard routing decision the Cluster Manager later "
    "uses to filter which jobs may run on which machines."
)

_body(doc, "Objective — minimize infrastructure cost; maximize fairness σ and workload diversity:")
_equation(doc, "min  λ₀ * infra_cost  −  λ₁ * σ  −  λ₂ * mix_total")

_body(doc, "The Cantelli second-order cone capacity constraint:")
_equation(doc, "C1a:  Σ_i f[i,n,h] + κ * t[n,h]  ≤  C_eff[n] * z[n,h]")
_equation(doc, "C1b:  Σ_i σ²[i,h] * y[i,n,h]  ≤  t[n,h]²    (SOCP cone)")

_body(doc,
    "Here κ = √((1 − ε) / ε) for chosen tail probability ε. "
    "Setting ε = 0.10 gives κ ≈ 3.0, providing a 90 percent guarantee that no node is "
    "overcommitted under the predicted demand variance σ²[i,h] from the prediction layer. "
    "This is the Cantelli bound, a standard chance-constraint technique for capacity planning "
    "under uncertainty."
)
_body(doc,
    "The plan-ahead output is a TenantAccessSchedule: for each planning period, the set of "
    "nodes each tenant is authorized to use. The real-time scheduler uses this as a filter "
    "restricting which jobs can be placed on which nodes each epoch. "
    "The dynamic modeling approach with time-indexed constraints follows Kovalenko and "
    "Zhdanova (2024), who formulate low-level Kubernetes scheduling as a combinatorial "
    "optimization problem with resource and activation constraints across time periods."
)

_heading(doc, "2.3  Iterative Decomposition Strategy", level=2)

_body(doc,
    "Both models grow intractably at production scale. The iterative strategy decomposes "
    "each problem into a sequence of smaller subproblems. This is a matheuristic in the sense "
    "of Maniezzo, Boschetti, and Stuetzle (2021): a heuristic built around an exact "
    "mathematical-programming core, derived from the classical divide-and-conquer decomposition "
    "principle in which a hard problem is split into smaller problems whose solutions are "
    "recomposed into a solution to the original. The applied precedent is Pandey and Patil "
    "(2022), who partition large-scale network problems into independent subnetworks and "
    "demonstrate computational savings of 15 to 68 percent over centralized solving. The same "
    "principle applies here: solve many small tractable instances rather than one large "
    "intractable one."
)

_body(doc, "Real-Time Iterative (Batch-MILP):")
_bullet(doc, "Sort unplaced jobs by predicted memory, largest first")
_bullet(doc, "Select the top batch of jobs and top batch of nodes by remaining capacity")
_bullet(doc, "Solve the sub-MILP to integer optimality")
_bullet(doc, "Update node usage, remove placed jobs, evict near-full nodes if spares exist")
_bullet(doc, "Repeat until all jobs are placed or the loop stalls")

_body(doc, "Plan-Ahead Iterative (Greedy FFD):")
_bullet(doc, "Maintain a sliding window of active tenants and active nodes")
_bullet(doc, "Sort tenants by remaining demand, largest first")
_bullet(doc, "For each tenant and period, allocate to the node with the most remaining capacity")
_bullet(doc, "Evict completed tenants and near-full nodes; provision fresh nodes as needed")
_bullet(doc, "Continue until all tenants are allocated")

_body(doc,
    "The real-time iterative solves each batch to integer optimality — it is a decomposition, "
    "not a relaxation. The plan-ahead iterative uses greedy first-fit decreasing, a "
    "polynomial-time algorithm that trades global optimality for solver-free speed at any scale. "
    "The mathematical formulations are unchanged; the iterative layer only avoids presenting "
    "the full problem to the solver at once."
)

_body(doc,
    "It is worth stating the scope of this claim precisely. The iterative strategy is "
    "theoretically aligned with decomposition: it follows the same divide-and-conquer logic that "
    "underlies established methods such as Dantzig-Wolfe, Lagrangian, and Benders decomposition, "
    "and it is a matheuristic in the sense of Maniezzo, Boschetti, and Stuetzle (2021). It is "
    "not, however, an implementation of any one of those formal methods. There is no master "
    "problem exchanging dual prices or cuts with subproblems and no convergence proof to the "
    "global optimum. The approach is a problem-specific decomposition heuristic inspired by the "
    "principle rather than a textbook algorithm applied verbatim, which is why it is reported as "
    "a matheuristic and not, for example, as Benders decomposition."
)

_page_break(doc)


# ── 2.4 IDEF-0 Model Representations ──────────────────────────────────────────
_heading(doc, "2.4  IDEF-0 Model Representations", level=2)
_body(doc,
    "The two models are summarized below as IDEF-0 boxes. Each box reads in the standard IDEF-0 "
    "convention: inputs enter from the left, constraints and controls enter from the top, the "
    "solvers and approach (mechanisms) enter from the bottom, and outputs leave to the right."
)
_image(doc, PLOTS / "idef_realtime.png", width=6.2,
       caption="Figure A — IDEF-0 representation of the Real-Time Scheduler (MILP)")
_image(doc, PLOTS / "idef_planahead.png", width=6.2,
       caption="Figure B — IDEF-0 representation of the Plan-Ahead Allocator (MISOCP)")

_page_break(doc)


# ── 3. Diagnostics and Model Evaluation ───────────────────────────────────────
_heading(doc, "3. Diagnostics and Model Evaluation")

_heading(doc, "3.1  Real-Time Model", level=2)

_heading(doc, "3.1.1  Non-Iterative: Computational Time Analysis", level=3)

_body(doc,
    "The single-shot MILP was benchmarked across a grid of job counts (J) and node counts (N) "
    "using four backends: Gurobi, CBC, SCIP, and HiGHS. Each solve was given a 60-minute "
    "wall-clock limit. Solve time grows super-linearly with J * N."
)
_table(doc,
    ["J", "N", "Variables", "Gurobi", "CBC", "Outcome"],
    [
        ["16",   "4",    "64",        "< 10 s",   "7 s",    "Optimal"],
        ["64",   "64",   "4,096",     "~15 s",    "~120 s", "Optimal"],
        ["256",  "256",  "65,536",    "~300 s",   "~600 s", "Time limit"],
        ["1024", "1024", "1,048,576", "Time limit","3,601 s","0% placed"],
    ]
)
doc.add_paragraph("")
_body(doc,
    "At J=1024, N=1024 both CBC and Gurobi exhaust the time limit. This confirms the "
    "exponential scaling of the binary program and motivates the iterative approach."
)

_heading(doc, "3.1.2  Iterative: Computational Time Analysis", level=3)

_body(doc,
    "The iterative RT solver was benchmarked at the largest scale point (J=1024, N=1024) across "
    "batch sizes, using Gurobi as the sub-MILP backend. Given an adequate per-sub-MILP solve "
    "budget, every batch size places 100 percent of jobs; solve time grows with batch size "
    "because each sub-MILP is exponentially harder, while the iteration count falls because each "
    "larger batch clears more jobs per solve."
)
_table(doc,
    ["J", "N", "Batch size", "Time", "Iterations", "Placement Rate"],
    [
        ["1024", "1024", "8 * 8",   "2.1 s",  "128", "100%"],
        ["1024", "1024", "16 * 16", "1.5 s",  "64",  "100%"],
        ["1024", "1024", "32 * 32", "15.2 s", "32",  "100%"],
        ["1024", "1024", "64 * 64", "32.2 s", "16",  "100%"],
    ]
)
doc.add_paragraph("")
_body(doc,
    "The batch=16*16 configuration achieves 100 percent placement in 1.5 seconds, roughly a 425x "
    "speedup over the single-shot Gurobi result of 635 seconds (which also places 100 percent but "
    "exhausts the 5-minute gap-closing limit). An earlier benchmark under a single shared time "
    "budget showed the large batches placing only a fraction of jobs; that was a budget artifact "
    "— each large sub-MILP could not return a solution within its small time slice and tripped "
    "the stall guard. With an adequate per-sub-MILP budget the placement rate is 100 percent at "
    "every batch size, confirming the cluster was never capacity-limited; small batches are "
    "preferred purely for speed."
)

_heading(doc, "3.1.3  Sensitivity Analysis", level=3)

_body(doc,
    "The sensitivity analysis sweeps K (SLA violation rolling window: 5 to 30) and "
    "jobs_per_round (arrival load: 10 to 100 jobs). Key findings:"
)
_bullet(doc, "K=10 provides the best balance: reacts to violations without over-penalizing node capacity")
_bullet(doc, "Placement rate drops from 100 percent at 10 jobs per round to below 40 percent at 100 jobs per round on a 5-node cluster")
_bullet(doc, "Queue growth becomes unbounded above the saturation threshold regardless of K; this is a capacity limitation, not a scheduling limitation")

_heading(doc, "3.1.4  Iterative vs Non-Iterative: Head-to-Head", level=3)

_body(doc,
    "The table below pairs the single-shot MILP against the iterative batch-MILP on the same "
    "(J, N) cells. Both columns use Gurobi, the production backend, so the comparison is "
    "apples-to-apples; the iterative solver uses a batch size of 16 * 16 (the fastest robust "
    "setting under Gurobi — see below). \"cap\" marks a run that hit the 5-minute (300 s) solve "
    "limit. Gurobi always returns a feasible 100-percent placement, but on the harder cells it "
    "exhausts the limit before closing the optimality gap, taking 5 to 11 minutes; the iterative "
    "solver reaches 100-percent placement in under two seconds. Iterations is the number of "
    "sub-MILP solves performed. Gap is the absolute wall-time saved and Gap % is that saving as "
    "a fraction of the single-shot time."
)
_table(doc,
    ["Jobs (J)", "Nodes (N)", "Batch size", "One-Shot Gurobi (s)",
     "Iterative Gurobi (s)", "Iterations", "Gap (s)", "Gap %"],
    [
        ["16",   "16",   "16 x 16", "1.1 (optimal)",   "0.03", "1",  "1.1",   "97.2%"],
        ["64",   "64",   "16 x 16", "21.4 (optimal)",  "0.33", "4",  "21.1",  "98.5%"],
        ["256",  "64",   "16 x 16", "304.2 (cap)",     "0.84", "16", "303.3", "99.7%"],
        ["256",  "256",  "16 x 16", "108.9 (optimal)", "0.57", "16", "108.4", "99.5%"],
        ["1024", "256",  "16 x 16", "547.6 (cap)",     "1.59", "64", "546.0", "99.7%"],
        ["1024", "512",  "16 x 16", "683.2 (cap)",     "1.25", "64", "681.9", "99.8%"],
        ["1024", "1024", "16 x 16", "635.2 (cap)",     "1.49", "64", "633.7", "99.8%"],
    ]
)
doc.add_paragraph("")
_body(doc,
    "Batch size is the one tuning knob. With each sub-MILP given an adequate solve budget, every "
    "batch size places 100 percent of jobs, but solve time grows sharply with batch size because "
    "each sub-MILP is exponentially harder. At J=1024, N=1024 the iterative solver (Gurobi) takes "
    "about 1.5 s at batch 16 * 16 (64 sub-solves), 2.1 s at 8 * 8 (128 sub-solves), 15 s at "
    "32 * 32 (32 sub-solves), and 32 s at 64 * 64 (16 sub-solves). Small batches win decisively: "
    "16 * 16 is the fastest because it balances few enough sub-solves against each one staying "
    "trivial for Gurobi, so it is adopted as the default (8 * 8 is essentially equivalent). "
    "(Under a tight shared time budget the large batches cannot solve any single sub-MILP in "
    "their slice and abandon the run early; this is a budget artifact, not a capacity limit, and "
    "is avoided by keeping the batch small.)"
)
_body(doc,
    "Jobs, not nodes, are the bottleneck. Reading across any row of the iterative heatmaps "
    "(fixed J, increasing N) the solve time is roughly constant — adding machines does not slow "
    "the solver, because solve time is driven by the number of jobs and therefore the number of "
    "sub-MILP iterations, not by the node count. This has a direct operational reading: if a "
    "10-second scheduling latency is the longest acceptable, the cluster can comfortably sustain "
    "well over a thousand pending jobs per round at a 16-job batch (1024 jobs place in about "
    "1.5 seconds), and the machine count is effectively free. This is the core justification for "
    "the 16-job default and for keeping the batch's node dimension small."
)
_body(doc,
    "The opposite regime — very few nodes — was tested separately (J = 8, 16 against "
    "N = 1, 2, 4, 8, 16, 32, both one-shot and iterative). There the limit is genuine capacity, "
    "not the solver: with a single node, 8 jobs place only about 38 percent and 16 jobs about 6 "
    "percent, because one machine cannot physically hold them; placement returns to 100 percent "
    "from N = 4 onward. Solve time across this small grid stays in the tens of milliseconds for "
    "both methods, confirming that at small scale the choice of solver or decomposition is "
    "immaterial — only physical capacity matters."
)

_heading(doc, "3.2  Plan-Ahead Model", level=2)

_heading(doc, "3.2.1  Non-Iterative: Computational Time Analysis", level=3)

_body(doc,
    "The plan-ahead MISOCP was benchmarked across tenant * node * period configurations "
    "with a 5-minute Gurobi time limit. Variable count grows as O(T * N * H) and build time "
    "grows quadratically with model size."
)
_table(doc,
    ["Tenants (T)", "Nodes (N)", "Variables", "Build (s)", "Solve (s)", "Status"],
    [
        ["8",   "4",   "< 500",     "0.1",   "1",    "Optimal"],
        ["32",  "64",  "~20k",      "5",     "10",   "Optimal"],
        ["128", "256", "241,793",   "491",   "301",  "Optimal"],
        ["128", "512", "483,585",   "553",   "301",  "Time Limit"],
        ["256", "256", "477,313",   "552",   "301",  "Time Limit"],
        ["256", "512", "954,625",   "816",   "301",  "Time Limit"],
        ["512", "512", "1,898,753", "1,035", "305",  "Time Limit"],
    ]
)
doc.add_paragraph("")
_body(doc,
    "The MISOCP finds an optimal solution at T=128, N=256 in approximately 8 minutes total. "
    "All larger configurations hit the time limit. At T=256, N=1024 the variable count "
    "reaches 1.9 million, causing out-of-memory errors on a 16 GB machine."
)

_heading(doc, "3.2.2  Iterative: Computational Time Analysis", level=3)

_body(doc,
    "The plan-ahead iterative (greedy FFD) was benchmarked across tenant counts where "
    "the full MISOCP fails. It uses no solver and runs in pure Python."
)
_table(doc,
    ["Tenants", "Iterations", "Satisfaction", "Elapsed"],
    [
        ["8",     "1",   "100%", "1 ms"],
        ["32",    "4",   "100%", "3 ms"],
        ["128",   "16",  "100%", "11 ms"],
        ["1,024", "~80", "100%", "79 ms"],
    ]
)
doc.add_paragraph("")
_body(doc,
    "The iterative approach places 128 tenants in 11 milliseconds versus 8 minutes for the "
    "full MISOCP, a speedup of over 43,000x. It achieves 100 percent demand satisfaction at "
    "all tested scales, including sizes where the MISOCP cannot even be built on 16 GB RAM."
)

_heading(doc, "3.2.3  Sensitivity Analysis", level=3)

_body(doc,
    "The plan-ahead sensitivity analysis sweeps Cantelli epsilon, exclusive tenant fraction, "
    "node capacity, and fairness weight λ₁. Key findings:"
)
_bullet(doc, "Smaller epsilon (stricter probabilistic guarantee) reduces effective node capacity by up to 30 percent, activating additional machines")
_bullet(doc, "Exclusive tenant fraction above 25 percent in small clusters risks model infeasibility")
_bullet(doc, "λ₁ > 2 yields diminishing fairness returns while substantially increasing solve time")
_bullet(doc, "MIP gap of 5 percent is the practical operating point; tightening to 1 percent multiplies solve time 5 to 10x with minimal improvement")

_heading(doc, "3.2.4  Iterative vs Non-Iterative: Head-to-Head", level=3)

_body(doc,
    "The table below pairs the single-shot MISOCP (Gurobi) against the iterative greedy first-fit "
    "allocator on the same (T, N) cells. The MISOCP column is total wall time (model build plus "
    "solve). \"cap\" marks a run whose solve phase hit the 5-minute (300 s) time limit; the total "
    "still exceeds five minutes because building the model alone takes minutes at these sizes. "
    "The iterative allocator uses no solver and achieves 100 percent demand satisfaction at every "
    "cell shown. Gap is the absolute wall-time saved and Gap % is that saving as a fraction of "
    "the single-shot time."
)
_table(doc,
    ["Tenants (T)", "Nodes (N)", "One-Shot MISOCP (s)", "Iterative FFD (s)",
     "Iterations", "Gap (s)", "Gap %"],
    [
        ["8",   "64",   "17.8 (optimal)",   "0.001", "1",  "17.8",    "99.997%"],
        ["8",   "1024", "66.7 (optimal)",   "0.005", "1",  "66.7",    "99.992%"],
        ["128", "256",  "338.3 (optimal)",  "0.022", "16", "338.3",   "99.994%"],
        ["128", "512",  "585.0 (cap)",      "0.038", "16", "585.0",   "99.994%"],
        ["256", "256",  "669.6 (cap)",      "0.038", "32", "669.6",   "99.994%"],
        ["256", "1024", "1,116.9 (cap)",    "0.132", "32", "1,116.8", "99.988%"],
        ["512", "512",  "1,340.1 (cap)",    "0.138", "64", "1,340.0", "99.990%"],
        ["512", "1024", "1,803.7 (cap)",    "0.258", "64", "1,803.5", "99.986%"],
    ]
)
_body(doc,
    "Iterations is the number of greedy windows the FFD allocator processed (the tenant "
    "population divided into windows of 8). The MISOCP solves the whole problem in one shot, so "
    "it has no analogous iteration count."
)
doc.add_paragraph("")

_page_break(doc)


# ── 4. Data Visualization ─────────────────────────────────────────────────────
_heading(doc, "4. Data Visualization and Communication")

_heading(doc, "4.1  Real-Time: Solve-Time Heatmaps (Non-Iterative)", level=2)
_body(doc,
    "Each cell shows median solve time for a (J, N) combination. Separate heatmaps are saved "
    "per backend. Dark cells indicate time-limit hits."
)
_image(doc, PLOTS / "rt_heatmap_CBC.png",   caption="Figure 1a — RT Solve-Time Heatmap (CBC)")
_image(doc, PLOTS / "rt_heatmap_HIGHS.png", caption="Figure 1b — RT Solve-Time Heatmap (HiGHS)")

_heading(doc, "4.2  Real-Time: Scaling (Non-Iterative)", level=2)
_image(doc, PLOTS / "rt_scaling_vars_vs_time.png",
       caption="Figure 2 — RT Scaling: Variable Count vs Solve Time (all backends)")

_heading(doc, "4.3  Real-Time: Iterative Heatmaps", level=2)
_body(doc, "Solve time heatmaps for the iterative solver at different batch sizes.")
_image(doc, PLOTS / "exp_rt_iter_heatmap_8x8.png",   caption="Figure 3a — RT Iterative (batch 8*8)")
_image(doc, PLOTS / "exp_rt_iter_heatmap_16x16.png", caption="Figure 3b — RT Iterative (batch 16*16)")

_heading(doc, "4.4  Real-Time: Iterative vs Non-Iterative Scaling", level=2)
_image(doc, PLOTS / "exp_rt_iter_scaling.png",
       caption="Figure 4 — RT Iterative vs Non-Iterative: Scaling Comparison")

_heading(doc, "4.5  Plan-Ahead: Solve-Time Heatmap (Non-Iterative)", level=2)
_image(doc, PLOTS / "pa_heatmap_solve_time.png",
       caption="Figure 5 — PA Solve-Time Heatmap (T * N grid, MISOCP)")
_image(doc, PLOTS / "pa_build_vs_solve_breakdown.png",
       caption="Figure 6 — PA Build vs Solve Time Breakdown")
_image(doc, PLOTS / "pa_vars_vs_solve_time.png",
       caption="Figure 7 — PA Variable Count vs Solve Time")

_heading(doc, "4.6  Plan-Ahead: Iterative Wall Time and Solve-Time Heatmap", level=2)
_image(doc, PLOTS / "exp_pa_iter_bar.png",
       caption="Figure 8 — PA Iterative: Wall Time vs Tenant Count (greedy FFD)")
_image(doc, PLOTS / "pa_iter_heatmap_solve_time.png",
       caption="Figure 8b — PA Iterative: Total Wall Time across the Tenant * Node grid "
               "(greedy FFD; mirrors the non-iterative MISOCP heatmap in Figure 5)")

_heading(doc, "4.7  Interactive Simulation Dashboard", level=2)
_body(doc,
    "A browser-based simulation dashboard (FastAPI backend, React frontend) visualizes the "
    "scheduling loop in real time. The dashboard shows node memory utilization bars, per-tenant "
    "wait times and fairness weights, a plan-ahead Gantt chart of tenant-node assignments per "
    "planning period, and batch-level statistics. All simulation parameters are configurable "
    "from the settings panel without restarting the backend."
)
_body(doc,
    "The dashboard includes an RT solver selector (Gurobi default), an iterative mode toggle, "
    "conditional batch size fields, and a Load Borg Config button that stages the Google "
    "cluster dataset parameters into the configuration in one click."
)
_image(doc, SCREENS / "dashboard.jpg", width=6.0,
       caption="Figure 9 — Simulation Dashboard: node cards, memory utilization wave, job queue, summary panel")
_image(doc, SCREENS / "planahead.jpg", width=6.0,
       caption="Figure 10 — Plan-Ahead Gantt Chart: tenant-machine assignments across 4 planning periods")

_heading(doc, "4.8  Iterative vs One-Shot — Bar Comparisons", level=2)
_body(doc,
    "The same head-to-head data as the tables in Sections 3.1.4 and 3.2.4, shown on a log scale "
    "so the difference is visible at a glance: the one-shot solver runs for minutes while the "
    "iterative method finishes in milliseconds to seconds."
)
_image(doc, PLOTS / "rt_headtohead_bar.png",
       caption="Figure 11 — Real-Time: one-shot Gurobi vs iterative (16 x 16), solve time (log scale)")
_image(doc, PLOTS / "pa_headtohead_bar.png",
       caption="Figure 12 — Plan-Ahead: one-shot MISOCP vs iterative greedy FFD, wall time (log scale)")

_heading(doc, "4.9  Small-Scale Stress Tests", level=2)
_body(doc,
    "The small-node RT grid (J = 8, 16 against N = 1..32) isolates the few-nodes regime, where "
    "the limit is physical capacity rather than the solver. The PA small grid (T = 8, 16, 32, 64 "
    "against N = 8, 16, 32, 64) confirms that iterative wall time tracks the tenant count and "
    "iterations, not the node count — each row is roughly flat."
)
_image(doc, PLOTS / "rt_small_grid_bar.png",
       caption="Figure 13 — RT small-node grid: one-shot vs iterative (J = 8, 16; N = 1..32)")
_image(doc, PLOTS / "rt_small_iter_heatmap.png",
       caption="Figure 14a — RT small-node grid (iterative 16 x 16): solve time, placement, iterations")
_image(doc, PLOTS / "rt_small_oneshot_heatmap.png",
       caption="Figure 14b — RT small-node grid (one-shot MILP): solve time and placement")
_image(doc, PLOTS / "pa_small_grid_heatmap.png",
       caption="Figure 15a — PA small grid (iterative FFD, no skips): wall time and iterations, T = 8..64, N = 8..64")
_image(doc, PLOTS / "pa_small_oneshot_heatmap.png",
       caption="Figure 15b — PA small grid (one-shot MISOCP, no skips): total wall time and status")

_page_break(doc)


# ── 5. Discussion and Implications ────────────────────────────────────────────
_heading(doc, "5. Discussion and Implications")

_heading(doc, "5.1  The Iterative Approach as Decomposition", level=2)

_body(doc,
    "The iterative strategy is a problem decomposition approach. Maniezzo, Boschetti, and "
    "Stuetzle (2021) describe this class of methods as matheuristics derived from "
    "decomposition: an exact solver is applied repeatedly to tractable subproblems within a "
    "heuristic master loop, following the same divide-and-conquer logic that underlies "
    "Dantzig-Wolfe, Lagrangian, and Benders decomposition. Pandey and Patil (2022) give an "
    "applied instance of the idea: their DSTAP-Heuristic partitions a road network into "
    "geographically separate subnetworks and achieves 15 to 68 percent computational savings "
    "over centralized optimization. The same principle applies here: the real-time batch-MILP "
    "partitions the J * N assignment problem into smaller windows solved to integer "
    "optimality, and the plan-ahead FFD loop partitions the tenant population into windows "
    "processed sequentially."
)
_body(doc,
    "For the real-time model each sub-MILP is solved to integer optimality, so the approach "
    "is a decomposition rather than a heuristic in the strict sense. For the plan-ahead model "
    "the greedy FFD allocation does not guarantee global optimality but achieves 100 percent "
    "demand satisfaction at all tested scales."
)

_heading(doc, "5.2  Placement Rate vs Consolidation Trade-Off", level=2)

_body(doc,
    "The iterative RT solver achieves a 100 percent placement rate but lower memory "
    "consolidation than the single-shot MILP. The global MILP maximizes the consolidation "
    "term, packing jobs onto fewer nodes. The iterative solver fills each batch greedily and "
    "leaves some nodes partially used. For operators prioritizing placement rate, the "
    "iterative approach is strictly preferable. For operators optimizing power usage by "
    "minimizing active node count, the single-shot MILP is better when tractable."
)

_heading(doc, "5.3  Prediction Layer Integration", level=2)

_body(doc,
    "The real-time model uses pred_mem_mb and pred_cpu_p95 directly in its constraints. "
    "The plan-ahead model uses u[i,h] and σ²[i,h] directly in the Cantelli capacity cone. "
    "The simulation dashboard can be switched to use live prediction API calls via the "
    "use_prediction_api configuration flag."
)
_body(doc,
    "One integration consideration: the prediction layer's u[i,h] values are normalized "
    "fractions (0 to 1 scale from the Google cluster traces). The plan-ahead model "
    "currently uses NODE_CAPACITY=10.0. For direct use of prediction API outputs, "
    "NODE_CAPACITY should be set to 1.0 in the Borg configuration, or the prediction layer "
    "should scale their u values by the chosen NODE_CAPACITY constant."
)

_page_break(doc)


# ── 6. Limitations ────────────────────────────────────────────────────────────
_heading(doc, "6. Limitations")

_bullet(doc,
    "Single resource dimension. The models optimize memory placement. CPU is enforced as "
    "a hard constraint but not optimized. Real workloads are constrained by both memory "
    "and CPU simultaneously."
)
_bullet(doc,
    "Plan-ahead capacity units. The MISOCP uses an abstract capacity unit that must be "
    "calibrated against the prediction layer's u[i,h] values. Mismatched units void the "
    "probabilistic guarantee."
)
_bullet(doc,
    "Greedy FFD sub-optimality. The plan-ahead iterative does not guarantee global "
    "optimality. Fairness sigma and the mix bonus objective are only available in the "
    "full Gurobi MISOCP."
)
_bullet(doc,
    "Batch size sensitivity. The iterative RT solver performance depends on batch "
    "configuration. Batch size must be tuned per deployment."
)
_bullet(doc,
    "No preemption or migration. The model is stateless per epoch. Running jobs cannot "
    "be moved between nodes."
)
_bullet(doc,
    "Synthetic workload data. All computational benchmarks use synthetically generated "
    "job and node profiles. Integration with actual prediction outputs is a pending step."
)

_page_break(doc)


# ── 7. Future Recommendations ─────────────────────────────────────────────────
_heading(doc, "7. Future Recommendations")

_heading(doc, "7.1  Near-Term", level=2)
_bullet(doc,
    "Integrate prediction API directly. Replace synthetic generators with prediction_api "
    "calls and validate that u[i,h] from the API produces feasible plan-ahead solutions "
    "with NODE_CAPACITY=1.0."
)
_bullet(doc,
    "Multi-resource optimization. Extend the real-time objective to include CPU using "
    "Dominant Resource Fairness so the placement objective balances both resources."
)
_bullet(doc,
    "Adaptive batch sizing. Tune RT batch sizes dynamically based on queue depth and "
    "node saturation rather than using fixed constants."
)

_heading(doc, "7.2  Medium-Term", level=2)
_body(doc,
    "Both models share the same block-angular structure — items (jobs or tenants) assigned to "
    "resources (nodes or machines), each resource carrying a coupling capacity limit — which is "
    "the textbook input for exact decomposition. A single method, Dantzig-Wolfe column "
    "generation (branch-and-price), therefore applies to both: a column is a feasible bundle on "
    "one resource (a memory-and-CPU knapsack for the real-time model, a Cantelli-cone packing for "
    "the plan-ahead model), the restricted master covers all items at optimal cost, and the "
    "pricing subproblem is solved per resource. Benders decomposition is the natural second fit, "
    "especially for the plan-ahead model, which already separates machine activation (master) "
    "from per-period allocation (subproblem). These exact methods would replace the current "
    "matheuristic with provable optimality at scale; the heuristic recomposition used here is a "
    "deliberate near-term shortcut, not an attempt at either method."
)
_bullet(doc,
    "Benders decomposition for the MISOCP. Following Maniezzo, Boschetti, and Stuetzle (2021), "
    "decompose the plan-ahead into a master problem (machine activation) and per-period "
    "subproblems, adding Benders cuts until the bound closes. This retains Gurobi optimality "
    "guarantees while scaling beyond T=1000 tenants."
)
_bullet(doc,
    "Column generation for the real-time MILP. Generate job-to-node columns on demand "
    "rather than enumerating all J * N, following standard techniques for large-scale "
    "bin-packing and assignment problems."
)

_heading(doc, "7.3  Long-Term", level=2)
_bullet(doc,
    "Online learning of Cantelli epsilon. Adapt the tail probability per node based on "
    "observed SLA violation rates, tightening bounds on stressed nodes and relaxing them "
    "on stable nodes to increase utilization."
)
_bullet(doc,
    "Preemption and migration. Extend the real-time model to support job migration across "
    "nodes, minimizing churn while improving global consolidation."
)

_page_break(doc)


# ── References ────────────────────────────────────────────────────────────────
_heading(doc, "References")

refs = [
    "Kovalenko, V., and Zhdanova, O. (2024). Dynamic Mathematical Model for Resource "
    "Management and Scheduling in Cloud Computing Environments. Information, Computing and "
    "Intelligent Systems, No. 5, 90-100. Igor Sikorsky Kyiv Polytechnic Institute. "
    "https://doi.org/10.20535.2786-8729.5.2024/316545",

    "Shi, Y., and Yu, H. (2024). Fairness-Aware Job Scheduling for Multi-Job Federated "
    "Learning. arXiv:2401.02740v3. Nanyang Technological University / Alibaba Group.",

    "Maniezzo, V., Boschetti, M. A., and Stuetzle, T. (2021). Decomposition based heuristics "
    "(Chapter 7). In Matheuristics: Algorithms and Implementations. EURO Advanced Tutorials on "
    "Operational Research. Springer, Cham. https://doi.org/10.1007/978-3-030-70277-9",

    "Pandey, V., and Patil, P. N. (2022). Computationally-Efficient Decomposition Heuristic "
    "for the Static Traffic Assignment Problem. arXiv:2206.12496v1. "
    "NC A&T State University / University of Texas at Austin.",

    "Verma, A., Pedrosa, L., Korupolu, M., Oppenheimer, D., Tune, E., and Wilkes, J. (2015). "
    "Large-scale cluster management at Google with Borg. ACM EuroSys.",

    "Google LLC. (2019). Google Cluster Usage Traces v3. "
    "https://github.com/google/cluster-data",
]
for r in refs:
    _bullet(doc, r)


# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f"\nReport saved: {OUT}")
