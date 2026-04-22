from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

NAVY  = RGBColor(0x0D, 0x1B, 0x2A)
BLUE  = RGBColor(0x00, 0x6A, 0xA6)
TEAL  = RGBColor(0x00, 0x7A, 0x8A)
DARK  = RGBColor(0x1A, 0x2E, 0x44)
GOLD  = RGBColor(0xB8, 0x86, 0x00)
GREEN = RGBColor(0x1E, 0x6B, 0x2E)
GRAY  = RGBColor(0x50, 0x50, 0x50)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY = RGBColor(0xF4, 0xF6, 0xF8)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def para(text, bold=False, italic=False, size=11, color=BLACK,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size); run.font.color.rgb = color
    return p

def heading(text, level=1):
    sizes  = {1: 15, 2: 12.5, 3: 11}
    colors = {1: BLUE, 2: TEAL, 3: DARK}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = colors.get(level, BLACK)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '0084AA')
        pBdr.append(bottom); pPr.append(pBdr)
    return p

def bullet(text, level=0, size=11, color=BLACK, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
        r.font.size = Pt(size); r.font.color.rgb = color
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color
    return p

def hline():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom); pPr.append(pBdr)

def labeled_para(label, text, label_color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(label + ":  "); r.bold = True
    r.font.size = Pt(11); r.font.color.rgb = label_color
    r2 = p.add_run(text); r2.font.size = Pt(11); r2.font.color.rgb = BLACK

def concept_entry(term, defn, term_color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(term + "  "); r.bold = True
    r.font.size = Pt(11); r.font.color.rgb = term_color
    r2 = p.add_run(defn); r2.font.size = Pt(11); r2.font.color.rgb = BLACK

# ══════════════════════════════════════════════════════════════════════════════
# COVER / TITLE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
run = p.add_run("Optimal Shared Memory Utilization with Service Level\nGuarantees in Multi-Tenant Clusters")
run.bold = True; run.font.size = Pt(18); run.font.color.rgb = NAVY

para("(Working Title — see Suggested Topic Sentences section for alternatives)",
     italic=True, size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para("DAMO 699 Capstone Project  ·  Master of Data Analytics",
     size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Supervisor: Hany Osman  ·  University of Niagara Falls  ·  Spring 2026",
     size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Group 3: Tha Pyay Hmu · Lhagii Tsogtbayar · Nadia Ríos · Jorge Mendoza · Alrick Grandison",
     size=10.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
hline()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — POINTERS
# ══════════════════════════════════════════════════════════════════════════════
heading("Pointers  (scope, platform, assumptions)")

labeled_para("Platform",
    "We focus on Kubernetes as our cluster management platform. It is the standard for "
    "multi-tenant workloads in production cloud systems, is well-documented, and has a formal "
    "Scheduling Framework that lets us plug in custom logic without modifying Kubernetes itself.")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(1)
r = p.add_run("Assumptions:"); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BLUE
bullet("(1) All workloads are containerized Kubernetes Jobs or Deployments — no VM-level scheduling.", level=1, size=10.5)
bullet("(2) Google Cluster Trace (optionally v3) is our primary dataset.", level=1, size=10.5)
bullet("(3) Simulation is the primary evaluation method — no live cluster required.", level=1, size=10.5)
bullet("(4) Memory is the primary scarce resource driving SLA violations — CPU throttling slows jobs down, memory OOM kills terminate them.", level=1, size=10.5)
bullet("(5) Tenants over-declare memory requests as a safety buffer — actual usage is consistently lower, and that gap is the overcommitment opportunity.", level=1, size=10.5)

labeled_para("Scope boundary",
    "We are not building a full cloud platform like AWS. We build the scheduling decision model "
    "(prediction + optimizer (prescription)) and evaluate it via simulation.")
labeled_para("Possible simplification",
    "Start with memory as the single constrained resource and CPU as secondary. This matches "
    "the group's original angle and keeps the math tractable for capstone scope.")
labeled_para("Additive peaks insight",
    "Model each job's memory usage as a pulse (ramps up, plateaus, ramps down). Total memory "
    "on a node = sum of all active pulses. Admission decision: will adding this new pulse cause "
    "the sum to cross the node's memory limit at any point? If yes → deny or delay.")
labeled_para("K8s hook points",
    "Our scheduler plugs into three Kubernetes phases: (1) Admission Webhook — ML-based "
    "admit/deny before scheduling; (2) Score phase — DRF fairness + temporal co-location "
    "ranking; (3) Runtime Controller — Prometheus-driven cgroup adjustments every 5 seconds.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SUGGESTED TOPIC SENTENCES
# ══════════════════════════════════════════════════════════════════════════════
heading("Suggested Topic Sentences")

options = [
    ("Option 1 — Preferred (captures the two-model pipeline and all three outcomes)",
     '"By using the output of a predictive model as the input of an optimization model that is '
     'constrained by available resources and SLA requirements, we can optimally schedule workloads '
     'on a cluster such that resource idleness is kept minimal and fairness to tenants is high to '
     'result in low operational cost and high client satisfaction."'),
    ("Option 2",
     '"This project designs an optimized predictive scheduling for multi-tenant Kubernetes clusters '
     'that closes the gap between declared resource requests and actual memory usage, enabling safe '
     'overcommitment while enforcing per-tenant SLA compliance and fairness."'),
    ("Option 3",
     '"We propose an optimization-based scheduling model for multi-tenant Kubernetes clusters that '
     'focuses on shared memory utilization as the primary constraint, using machine learning to '
     'dynamically predict resource requirements, maximize cluster utilization, minimize SLA '
     'violations, and maintain fairness across tenants."'),
    ("Option 4",
     '"Cloud clusters waste up to 40\u201360% of available memory because schedulers rely on declared '
     'requests rather than actual usage \u2014 this project builds a smarter Kubernetes scheduler that uses '
     'tenant workload history to predict real memory demand, safely pack more jobs onto each node, '
     'and prevent SLA violations before they occur."'),
]
for label, text in options:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + "  ")
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = GREEN if "Preferred" in label else DARK
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.paragraph_format.space_before = Pt(1); p2.paragraph_format.space_after = Pt(5)
    r2 = p2.add_run(text); r2.italic = True; r2.font.size = Pt(11); r2.font.color.rgb = DARK

para("(Working Topic Sentence)", italic=True, size=10, color=GRAY, space_after=2)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — SUGGESTED HYPOTHESIS
# ══════════════════════════════════════════════════════════════════════════════
heading("Suggested Hypothesis")
para(
    "We hypothesize that a Kubernetes scheduler augmented with (1) machine learning-based "
    "prescriptive scheduling, and an optimization model that incorporates (2) Dominant Resource "
    "Fairness-aware placement, and (3) dynamic cgroup-based SLA enforcement will achieve "
    "significantly higher cluster memory utilization than the default Kubernetes scheduler while "
    "maintaining SLA compliance rates above 95% and reducing inter-tenant fairness inequality, "
    "as measured by Gini coefficient and dominant share variance, on workloads replayed from "
    "the Google Cluster Trace v3.",
    space_after=6
)
para("Specifically, we expect:", bold=True, size=11, space_after=2)
bullet("Cluster memory utilization to increase from ~45–60% (default K8s) to \u2265 85%")
bullet("SLA violation rate to remain below 5% — matching or improving on Priya (2025)'s benchmark")
bullet("Gini coefficient across tenants to decrease from ~0.25 toward ~0.10 — matching Alatawi (2025)'s range")
bullet("Prediction accuracy (MAPE) below 5% using Random Forest on tenant history features")
para("(Working Hypothesis — is it required?)", italic=True, size=10, color=GRAY, space_after=2)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — SUGGESTED INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading("Suggested Introduction")
for text in [
    "Cloud infrastructure costs are escalating on multiple fronts. Conventional DRAM supplies are "
    "contracting as manufacturers redirect production toward high-bandwidth memory for AI accelerators, "
    "with prices projected to rise 54\u2013116% year-over-year (S&P Global, 2026). The operational cost of "
    "leaving compute and memory resources idle compounds these pressures. This creates a compelling "
    "imperative for cloud providers: extract significantly more value from existing hardware before "
    "buying more. The bottleneck is not hardware availability \u2014 it is the scheduler.",

    "Current cluster memory utilization in production environments averages between 40% and 60%, meaning "
    "that on any given node, nearly half its memory sits idle while tenants wait for their jobs to start "
    "(Chaudhari, 2025). This waste exists not because hardware is unavailable, but because schedulers are "
    "conservative. The default Kubernetes scheduler places workloads based solely on declared resource "
    "requests and current availability, with no awareness of how much memory jobs will actually use, when "
    "co-located jobs will peak, or whether admitting a new job will cause a neighbor\u2019s SLA to be violated "
    "in the next thirty minutes. The result is predictable: either the scheduler over-admits and memory "
    "contention causes OOM kills, or it under-admits and expensive hardware idles.",

    "Multi-tenancy sharpens the problem. When multiple teams share the same cluster, a single greedy "
    "tenant can monopolize memory, causing others to queue arbitrarily long. Kubernetes has no native "
    "fairness mechanism \u2014 it processes requests in arrival order without regard to whether one tenant has "
    "already consumed a disproportionate share of cluster memory. Latency-sensitive jobs compete against "
    "batch workloads with no systematic prioritization, and memory violations are discovered after the "
    "fact rather than prevented proactively.",

    "This project addresses the gap between what the Kubernetes scheduler currently does and what a "
    "production multi-tenant cluster requires. The gap is precise: no existing system combines "
    "prescriptive scheduling, multi-resource fairness, and runtime SLA enforcement into a single coherent "
    "Kubernetes scheduler. Chaudhari (2025) describes the need for exactly this combination \u2014 a Workload "
    "Classifier, Fairness Engine, and Predictive Resource Manager \u2014 but the framework exists only as a "
    "proposal with no implementation or evaluation. We build it.",

    "The proposed system is a custom Kubernetes scheduler plugin with two integrated analytical "
    "components: a Predictive Model and an Optimization Model (Prescriptive). The Predictive Model "
    "trains a Random Forest on tenant workload history from the Google Cluster Trace to predict the "
    "actual peak memory utilization of each submitted job. The Optimization Model then takes those "
    "predictions as inputs and decides whether and where to place each job, applying Dominant Resource "
    "Fairness (Ghodsi et al., 2011) to ensure no tenant monopolizes memory, and enforcing SLA deadline "
    "constraints before any placement is confirmed. A runtime enforcement layer assigns memory priority "
    "weights to pods via Kubernetes QoS classes and dynamic cgroup controls.",

    "The system is evaluated against the Google Cluster Trace v3 using discrete-event simulation, "
    "comparing against the default Kubernetes scheduler and a static DRF baseline. Target outcomes are "
    "cluster memory utilization above 85%, SLA compliance above 95%, and inter-tenant fairness variance "
    "below 10%.",
]:
    para(text, space_after=8)

para("(Working — to be made more concise)", italic=True, size=10, color=GRAY, space_after=2)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — OPTIMIZATION MODEL DESIGN
# ══════════════════════════════════════════════════════════════════════════════
heading("Suggested Optimization Model Design")
para("Adapting Kovalenko & Zhdanova (2024) as the structural foundation, extended with DRF fairness, "
     "Random Forest prediction, and SLA deadline constraints.",
     italic=True, size=10.5, color=GRAY, space_after=6)
para(
    "Kovalenko & Zhdanova (2024) provide one of the few Kubernetes scheduling papers that writes down "
    "actual mathematical constraints \u2014 typed variables, objectives, and hard bounds \u2014 rather than "
    "describing goals in prose. Their discrete combinatorial model assigns pods to nodes across a time "
    "horizon and is the closest existing formulation to what our capstone needs. We take it as the "
    "structural skeleton and make four targeted extensions.",
    space_after=6
)

for label, items in [
    ("What we keep from Kovalenko", [
        "Binary pod-to-node assignment variable: x_{n,j} \u2208 {0,1}",
        "Per-node resource capacity constraints (CPU and memory per time period)",
        "Multi-objective structure (utilization + penalty terms)",
        "Discrete time horizon T (each period = ~5-minute snapshot)",
    ]),
    ("What we change", [
        "Declared requests \u2192 predicted utilization: replace static cp, mp with RF-predicted \u0109_j(t), m\u0302_j(t). "
        "This is the core enabler of safe overcommitment \u2014 scheduling on predictions rather than inflated declarations.",
        "Add DRF fairness constraint: d_k = max(memory share, CPU share) per tenant; constrain d_k \u2264 (1/|K|) + \u03b5.",
        "Add SLA deadline constraint: T_start + dur_j \u2264 D_j. Jobs that cannot meet their deadline are queued.",
        "Add temporal peak-time safety check: enforce memory capacity at predicted peak time t*_j only.",
        "Remove server power on/off objective: not applicable to pre-provisioned Kubernetes nodes.",
    ]),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = TEAL
    for item in items:
        bullet(item, level=0, size=10.5)

para("Condensed objective function:", bold=True, size=11, space_before=6, space_after=2)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.35)
p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
r = p.add_run("maximize:  w\u2081 \u00b7 (avg memory utilization)  \u2212  w\u2082 \u00b7 (SLA violation penalty)  \u2212  w\u2083 \u00b7 (dominant share spread)")
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = DARK

para("Suggested starting weights: w\u2081 = 0.5 (utilization), w\u2082 = 0.35 (SLA protection), w\u2083 = 0.15 (fairness).",
     size=10.5, color=GRAY, space_after=4)
para("For capstone simulation: implement as a greedy sequential admission algorithm \u2014 check constraints "
     "in order (memory safety at predicted peak \u2192 CPU safety \u2192 fairness bound \u2192 SLA deadline) \u2192 if all pass, "
     "admit; otherwise, queue. This is tractable at simulation scale and matches how production schedulers work.",
     size=10.5, space_after=4)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — PREDICTION MODEL DESIGN
# ══════════════════════════════════════════════════════════════════════════════
heading("Suggested Prediction Model Design")
para("Built on Resource Central (Cortez 2017) as the architecture, Random Forest justified by Doukha "
     "(2025), preprocessing from Kofi (2025).",
     italic=True, size=10.5, color=GRAY, space_after=6)
para("The optimization model above needs m\u0302_j(t) and \u0109_j(t) \u2014 predicted memory and CPU usage of each job "
     "at each point in time. The prediction model computes these before the scheduler makes its placement decision.",
     space_after=6)

para("Three things to predict:", bold=True, size=11, space_after=2)
bullet("P95 memory peak m\u0302_j^peak \u2014 worst-case usage for the optimization model\u2019s memory safety check")
bullet("Temporal usage profile m\u0302_j(t) \u2014 the shape of usage over the job\u2019s lifetime, used for co-location scoring")
bullet("Job duration dur_j \u2014 when resources free up, used for the SLA deadline check")

para("Key features (Random Forest inputs):", bold=True, size=11, space_before=6, space_after=2)
for feat in [
    "Tenant\u2019s historical actual/declared memory ratio \u2014 the single most predictive feature",
    "Tenant\u2019s historical actual/declared CPU ratio",
    "Declared memory request and CPU request (job-level)",
    "Job priority class (BestEffort / Burstable / Guaranteed)",
    "Hour of day and day of week (temporal usage cycle features)",
    "Current node memory utilization",
    "Hours until the peak of the largest job currently running on the candidate node",
]:
    bullet(feat, size=10.5)

para("Model:", bold=True, size=11, space_before=6, space_after=2)
para("Random Forest Regressor (Doukha 2025: MAPE 2.65% vs LSTM\u2019s 17.43%). Preprocessing: "
     "Savitzky-Golay smoothing + min-max normalization on Google Cluster Trace v3 (Kofi 2025 pipeline). "
     "Output scaled by a 5% uncertainty buffer before being fed into the optimization model.",
     size=10.5, space_after=4)
para("Temporal profile: trapezoid pulse \u2014 ramp up 10% of duration, plateau 80%, ramp down 10%. "
     "Total node memory at time t = sum of all active pulses. Decision check: does adding this new pulse "
     "cause the sum to exceed \u03b1 \u00b7 RAM_n at any point?",
     size=10.5, space_after=4)
para("Drift handling: retrain weekly or when prediction MAPE on recent jobs exceeds 8% (Perera 2025).",
     size=10.5, space_after=4)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — BRIEF LITERATURE REVIEW TABLE
# ══════════════════════════════════════════════════════════════════════════════
heading("Brief Literature Review")
para("21 papers reviewed. Table below summarizes each paper's contribution, the gap it addresses, "
     "what we take from it, and what kind of model or algorithm it uses.",
     size=10.5, color=GRAY, italic=True, space_after=8)

col_widths = [Inches(1.45), Inches(1.65), Inches(1.25), Inches(1.65), Inches(1.85)]
col_headers = ["Paper", "Core Contribution", "Gap Addressed", "What We Include", "Model Type"]

tbl = doc.add_table(rows=1 + 21, cols=5)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_row = tbl.rows[0]
for cell, w, hdr in zip(hdr_row.cells, col_widths, col_headers):
    cell.width = w
    shade_cell(cell, '1A2E44')
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(hdr)
    run.bold = True; run.font.size = Pt(9); run.font.color.rgb = WHITE

rows_data = [
    ("Chaudhari (2025)",
     "K8s fails for AI/ML; proposes 4-component framework \u2014 never built",
     "Gap anchor",
     "Architecture blueprint; util. baselines",
     "None \u2014 conceptual framework"),
    ("Jiang Zhi (2025)",
     "Overcommitment study; Clovers Python simulator; real trace datasets",
     "Simulation environment",
     "Simulator design + trace datasets",
     "None \u2014 rule-based simulator"),
    ("Coach / Reidys (2025)",
     "Co-locate workloads with non-overlapping temporal memory peaks; +26% capacity",
     "Overcommitment strategy",
     "Temporal co-location scoring",
     "Prediction: temporal pattern forecasting"),
    ("DRF \u2014 Ghodsi (2011)",
     "Dominant Resource Fairness: LP-proven multi-resource fair allocation",
     "Fairness backbone",
     "DRF algorithm as Score phase",
     "Fairness algorithm: LP-proven dominance"),
    ("Resource Central (2017)",
     "RF predicts P95 CPU/memory from tenant history; 79\u201390% accuracy",
     "Prediction architecture",
     "RF model + tenant history features",
     "Prediction model: Random Forest (P95)"),
    ("Blagodurov et al.",
     "cgroups weights for critical vs batch; near-100% util. with SLA maintained",
     "Runtime SLA enforcement",
     "Dynamic cgroup weight controller",
     "Scheduler optimizer: cgroups weight equations"),
    ("Wang & Yang (2025)",
     "LSTM+DQN on Kubernetes; 32.5% util. gain, 43.3% latency reduction",
     "K8s deployment blueprint",
     "Plugin architecture; comparison baseline",
     "Prediction (LSTM) + RL scheduler (DQN)"),
    ("Doukha (2025)",
     "RF beats LSTM: MAPE 2.65% vs 17.43% on CPU/memory prediction",
     "Model selection proof",
     "RF as primary predictor; MAPE metric",
     "Prediction comparison: RF vs LSTM"),
    ("Atropos / Hu (2025)",
     "Throttle culprit workload during overload \u2014 not innocent victim",
     "Overload fallback",
     "Targeted pod throttling on overload",
     "None \u2014 monitoring-based detection"),
    ("Priya (2025)",
     "SloPolicy CRD + K8s plugin + Prometheus loop; 45% P99 reduction, <5% SLA",
     "K8s architecture reference",
     "CRD design; plugin structure; eval metrics",
     "Scheduler optimizer: QoS scoring function"),
    ("Kofi (2025)",
     "LSTM on Google Cluster Trace v3; R\u00b2=0.99 with preprocessing pipeline",
     "Dataset + preprocessing",
     "Google Trace + Savitzky-Golay pipeline",
     "Prediction model: LSTM (R\u00b2=0.99)"),
    ("Perera (2025/2026)",
     "RL schedulers: model drift and interpretability issues in production",
     "Model choice justification",
     "RF interpretability arg.; drift detection",
     "None \u2014 review / landscape paper"),
    ("Pinnapareddy (2025)",
     "Right-sizing, bin packing, autoscaling in K8s; cost framing",
     "Motivation + tooling",
     "Kubecost for per-tenant eval.",
     "None \u2014 practitioner analysis"),
    ("Patchamatla",
     "K8s on OpenStack; VM-hosted vs bare-metal; scheduler coordination",
     "Deployment context",
     "Architecture reference",
     "None \u2014 experimental comparison"),
    ("Liu & Guitart (2025)",
     "In-node DRC: group-aware cgroup tuning; 242\u2013319% throughput gain",
     "In-node enforcement",
     "DRC concept post-placement",
     "Scheduler optimizer: in-node cgroup"),
    ("Kovalenko (2024)",
     "Formal discrete combinatorial optimization model for K8s scheduling",
     "Math formalization",
     "Constraint + objective structure",
     "Scheduler optimizer: discrete comb. LP"),
    ("Alatawi (2025)",
     "RL MDP for serverless multitenancy; Gini fairness; 98% SLA, 50% latency",
     "Fairness metric + RL compare",
     "Gini coefficient metric; RL baseline",
     "RL policy model: MDP-based allocation"),
    ("Zhao et al. (2021)",
     "Formal admission control + profit optimization for AaaS; deadline + budget SLA",
     "Optimization model structure",
     "Admission algorithm + SLA constraints",
     "Scheduler optimizer: admission LP"),
    ("Borg \u2014 Verma (2015)",
     "Google large-scale cluster mgmt; overcommit via sharing; 20\u201330% more machines without sharing",
     "Foundational prior art",
     "Overcommit justification at scale",
     "None \u2014 engineering framework"),
    ("Quasar \u2014 Delimitrou (2014)",
     "QoS-aware scheduling; ML profiling replaces static reservation; <20% CPU with static",
     "Underutilization validation",
     "QoS-aware framing; profiling reference",
     "Prediction: collaborative filtering"),
    ("Heracles \u2014 Lo (2015)",
     "Memory bandwidth is primary co-location bottleneck; ~90% util. + <5% SLA with feedback ctrl",
     "Memory-first justification",
     "Memory contention evidence",
     "Scheduler optimizer: feedback-based ctrl"),
]

for j, row_data in enumerate(rows_data):
    row = tbl.rows[j + 1]
    bg = 'F4F6F8' if j % 2 == 0 else 'FFFFFF'
    for i, (cell, val, w) in enumerate(zip(row.cells, row_data, col_widths)):
        cell.width = w
        shade_cell(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        if i == 0:
            run.bold = True; run.font.color.rgb = TEAL
        elif i == 4:
            run.font.color.rgb = BLUE
        else:
            run.font.color.rgb = BLACK

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# NOTE ON TEAM DRAFT
# ══════════════════════════════════════════════════════════════════════════════
hline()
heading("Note on Team Draft Contributions", level=2)
para(
    "The team's initial proposal draft established several important framing decisions: "
    "(1) memory as the primary constraint in a multi-tenant cluster setting; (2) the scope "
    "boundary \u2014 not building AWS but a scheduling model and simulation; (3) the core objective "
    "tuple: maximize utilization, minimize SLA violations, maintain fairness; and (4) mathematical "
    "optimization with real cluster data (Google Cluster Trace) as the analytical approach.",
    size=10.5, color=GRAY, italic=True, space_after=8
)
hline()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — BASIC IDEA
# ══════════════════════════════════════════════════════════════════════════════
heading("Basic Idea  (simple explanation for the team)")
para(
    "Imagine a shared office building where multiple companies rent desk space. Each company tells "
    "the building manager 'I need 20 desks,' but on any given day they actually use 12. The building "
    "manager, following the rules, keeps 20 desks reserved per company \u2014 so half the building sits "
    "empty even when other companies are waiting for a single desk. To make things worse, some "
    "companies have urgent meetings (latency-sensitive jobs) while others are doing background filing "
    "(batch jobs), but they all wait in the same queue with no priority distinction.",
    space_after=6
)
para(
    "Now imagine a smarter building manager who: (1) looks at how much desk space each company has "
    "historically used \u2014 not just what they asked for \u2014 and reserves only what they will likely "
    "actually need; (2) gives priority to the company that has been waiting longest or has used the "
    "least space so far; and (3) guarantees that companies with urgent meetings always get their "
    "desks on time, even if the filing team has to wait.",
    space_after=6
)
para(
    "That is what this project builds \u2014 but for cloud computing. The 'building' is a Kubernetes "
    "cluster. The 'companies' are tenants. The 'desks' are memory (and CPU). The smarter building "
    "manager is our custom Kubernetes scheduler. We test it by replaying real workload data from "
    "Google's production cluster through a simulation \u2014 no live cloud account needed.",
    space_after=4
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — KEY CONCEPTS
# ══════════════════════════════════════════════════════════════════════════════
heading("Key Concepts")
para("Definitions for team members who are new to cluster scheduling.",
     size=10.5, color=GRAY, italic=True, space_after=6)

concept_entry("Multi-Tenant Cluster",
    "A shared pool of servers where multiple organizations ('tenants') submit workloads. "
    "They share the same physical resources but expect isolation and fairness. "
    "Example: a university research cluster shared by several research groups.")

concept_entry("Kubernetes",
    "The dominant open-source cluster manager \u2014 same category as Google Borg, Apache Mesos. "
    "It accepts workload submissions, decides which server runs each one, and keeps them running. "
    "Think of it as the operating system for a cluster.")

concept_entry("Pod / Container",
    "The unit of work in Kubernetes. A container holds one application and its dependencies, "
    "starts in seconds, and shares the host OS kernel. A pod is one or more containers scheduled "
    "together. When we say 'job,' we mean a pod submitted by a tenant.")

concept_entry("Scheduler vs. Optimizer \u2014 what is the difference?",
    "The SCHEDULER is the mechanism: the Kubernetes component that assigns pods to nodes. "
    "The OPTIMIZER is the decision logic: the mathematical/ML model that tells the scheduler "
    "which choice is best. In our project, we build the optimizer and plug it into Kubernetes "
    "as a scheduler plugin.")

concept_entry("Predictive Model vs. Optimization Model (Prescriptive) \u2014 the two things we build",
    "The Predictive Model (Random Forest) answers: 'How much memory will this job actually use, "
    "and when?' The Optimization Model (Prescriptive) takes those predictions and answers: "
    "'Should we admit this job right now, and if so, on which node?' It applies constraints "
    "(memory capacity, fairness bounds, SLA deadlines). In data analytics terms: predictive "
    "analytics \u2192 prescriptive analytics (also called operations research / decision optimization).")

concept_entry("Admission Control",
    "The front-door decision point: 'Should we accept this job right now?' This is what the "
    "Optimization Model executes. It takes the Predictive Model's output and applies all "
    "constraints to decide: admit, queue, or deny. In our project, our Optimization Model IS "
    "the admission control logic.")

concept_entry("Resource Overcommitment",
    "Accepting more workloads than the server's declared capacity on paper, relying on the fact "
    "that not all jobs will peak simultaneously. Safe when done with accurate prediction. "
    "Dangerous without it.")

concept_entry("SLA (Service Level Agreement)",
    "A formal contract stating what level of service a tenant is guaranteed \u2014 e.g., 'your job "
    "completes within X minutes.' Violating the SLA has financial consequences. Our system "
    "prevents violations proactively, not reactively.")

concept_entry("Memory OOM Kill",
    "When a container exceeds its memory limit, the Linux kernel kills it immediately \u2014 no "
    "warning. This is fundamentally different from CPU throttling, which just slows the job "
    "down. Memory violations are fatal and directly violate SLAs.")

concept_entry("Dominant Resource Fairness (DRF)",
    "A fairness algorithm from Ghodsi et al. (2011). For each tenant, find which resource "
    "(CPU or memory) they consume the largest fraction of across the cluster. DRF ensures no "
    "tenant's dominant share grows disproportionately. Proven: no tenant prefers another's "
    "allocation, everyone gets their fair share, cannot be gamed.")

concept_entry("Gini Coefficient",
    "A number between 0 and 1 measuring inequality in resource distribution. 0 = perfect "
    "equality. Alatawi (2025) reduces it from 0.25 to 0.10 \u2014 that improvement range is our target.")

concept_entry("Google Cluster Trace v3",
    "A publicly available dataset from Google: real job records with arrival times, declared "
    "requests, actual measured usage, duration, tenant IDs, priority levels. "
    "This is our training and simulation dataset.")

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = r"C:\Users\alric\Documents\Spring2026\Capstone_ClaudeAI\Proposal_Draft.docx"
doc.save(out)
print(f"Saved: {out}")
