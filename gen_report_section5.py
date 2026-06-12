"""
gen_report_section5.py
──────────────────────
Produces a copy of optimization_layer_report.docx for pasting into the main
combined document. It does NOT change any wording. It only:

  • drops the title block and Section 1 (Introduction / Abstract),
  • starts the content at "Analytical Methods",
  • renumbers every top-level section by +3 (2 -> 5, 3 -> 6, ... 7 -> 10),
    so the first section becomes "5. Analytical Methods",
  • sets all heading colors to normal black (no theme blue).

It reads the existing optimization_layer_report.docx and writes a SEPARATE
file. The original document is never modified.

Run:  python gen_report_section5.py
Output: optimization_layer_report_section5.docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
SRC  = ROOT / "optimization_layer_report.docx"
OUT  = ROOT / "optimization_layer_report_section5.docx"

# ── Configurable ───────────────────────────────────────────────────────────────
# The first kept section ("Analytical Methods") is renumbered to this value, and
# every following section follows on (6, 7, ...). Change this one number to slot
# the content in after a different section of your main document.
START_SECTION_NUMBER = 5

# Where figure and table numbering starts. The first table becomes "Table N" and
# each following table counts up; figures are offset the same way (the first
# figure group becomes "Figure N", keeping any a/b suffixes). Change as needed.
TABLE_START_NUMBER  = 5
FIGURE_START_NUMBER = 5

# APA titles for the four data tables, in document order. Edit if wording changes.
TABLE_TITLES = [
    "Real-Time Single-Shot MILP Solve Time by Problem Size",
    "Real-Time Iterative Solver Performance by Batch Size",
    "Plan-Ahead MISOCP Solve Time by Problem Size",
    "Plan-Ahead Iterative (Greedy FFD) Performance by Tenant Count",
]

# ── Internal ─────────────────────────────────────────────────────────────────
_FIRST_SECTION_SOURCE_NUMBER = 2                       # "2. Analytical Methods"
_FIRST_FIGURE_SOURCE_NUMBER  = 1                       # first caption is "Figure 1.."
START_HEADING    = f"{_FIRST_SECTION_SOURCE_NUMBER}. Analytical Methods"
OFFSET           = START_SECTION_NUMBER - _FIRST_SECTION_SOURCE_NUMBER
FIGURE_OFFSET    = FIGURE_START_NUMBER - _FIRST_FIGURE_SOURCE_NUMBER
FIGURE_SEPARATOR = "—"                            # em dash in existing captions
BLACK            = RGBColor(0x00, 0x00, 0x00)


def _drop_front_matter(doc: Document) -> None:
    """Remove every body element before the Analytical Methods heading."""
    body   = doc.element.body
    target = None
    for p in doc.paragraphs:
        if p.text.strip().startswith(START_HEADING):
            target = p._element
            break
    if target is None:
        raise RuntimeError(f"Could not find heading starting with {START_HEADING!r}")

    for child in list(body):
        if child is target:
            break
        if child.tag == qn("w:sectPr"):   # keep the section properties block
            continue
        body.remove(child)


def _renumber_and_recolor(doc: Document) -> None:
    """Bump the leading section number by OFFSET and force black heading text."""
    def _bump(match: re.Match) -> str:
        return str(int(match.group(1)) + OFFSET)

    for p in doc.paragraphs:
        style = p.style.name or ""
        if not style.startswith("Heading"):
            continue
        # Renumber: leading integer (handles "2", "2.1", "3.1.1"); References has none.
        if p.runs:
            new_text = re.sub(r"^(\d+)", _bump, p.runs[0].text)
            if new_text != p.runs[0].text:
                p.runs[0].text = new_text
        # Recolor every run in the heading to normal black.
        for run in p.runs:
            run.font.color.rgb = BLACK


def _make_label(doc, text, *, bold=False, italic=False):
    """A flush-left label paragraph (APA figure/table number or title)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = BLACK
    return p


def _apa_figures(doc) -> None:
    """
    APA 7 figure layout: move the caption ABOVE the image as two flush-left
    lines — 'Figure N' (bold) then the title (italic). The original wording of
    each caption is preserved; only its placement and styling change.
    """
    paras = doc.paragraphs
    items = []
    for i, p in enumerate(paras):
        if p._element.xpath(".//w:drawing"):
            cap = paras[i + 1] if i + 1 < len(paras) else None
            if cap is not None and cap.text.strip().lower().startswith("figure"):
                items.append((p, cap))

    for img_p, cap in items:
        text = cap.text.strip()
        if FIGURE_SEPARATOR in text:
            label, title = (s.strip() for s in text.split(FIGURE_SEPARATOR, 1))
        else:
            label, title = text, ""
        # Offset the figure's leading number, keeping any a/b suffix (e.g. "1a").
        label = re.sub(r"^(Figure\s+)(\d+)",
                       lambda m: f"{m.group(1)}{int(m.group(2)) + FIGURE_OFFSET}",
                       label)
        num_p = _make_label(doc, label, bold=True)
        img_p._element.addprevious(num_p._element)          # -> [num, image]
        if title:
            title_p = _make_label(doc, title, italic=True)
            img_p._element.addprevious(title_p._element)     # -> [num, title, image]
        cap._element.getparent().remove(cap._element)        # drop old below-caption


def _apa_tables(doc) -> None:
    """
    APA 7 table layout: add 'Table N' (bold) and an italic title ABOVE each
    table, flush left. Titles come from TABLE_TITLES (editable above).
    """
    for idx, table in enumerate(doc.tables):
        n     = TABLE_START_NUMBER + idx
        title = TABLE_TITLES[idx] if idx < len(TABLE_TITLES) else f"Table {n}"
        num_p   = _make_label(doc, f"Table {n}", bold=True)
        title_p = _make_label(doc, title, italic=True)
        table._element.addprevious(num_p._element)            # -> [num, table]
        num_p._element.addnext(title_p._element)              # -> [num, title, table]


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Source not found: {SRC} (run gen_report.py first)")
    doc = Document(str(SRC))          # read-only; original is never written
    _drop_front_matter(doc)
    _renumber_and_recolor(doc)
    _apa_figures(doc)
    _apa_tables(doc)
    doc.save(str(OUT))
    print(f"Section-5 report saved (original untouched): {OUT}")


if __name__ == "__main__":
    main()
