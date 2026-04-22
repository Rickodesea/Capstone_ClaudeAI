from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

NAVY  = RGBColor(0x0D, 0x1B, 0x2A)
BLUE  = RGBColor(0x00, 0x6A, 0xA6)
TEAL  = RGBColor(0x00, 0x7A, 0x8A)
DARK  = RGBColor(0x1A, 0x2E, 0x44)
GOLD  = RGBColor(0xB8, 0x86, 0x00)
GREEN = RGBColor(0x1E, 0x6B, 0x2E)
GRAY  = RGBColor(0x50, 0x50, 0x50)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED   = RGBColor(0x8B, 0x00, 0x00)

MEMBER_COLORS = [
    RGBColor(0x00, 0x5F, 0x73),  # Member 1 — teal-blue
    RGBColor(0x94, 0x00, 0x56),  # Member 2 — deep magenta
    RGBColor(0x1E, 0x6B, 0x2E),  # Member 3 — dark green
    RGBColor(0xB8, 0x60, 0x00),  # Member 4 — amber
    RGBColor(0x3D, 0x00, 0x8F),  # Member 5 — purple
]

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def para(text, bold=False, italic=False, size=11, color=BLACK,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size); run.font.color.rgb = color
    return p

def heading(text, level=1, color=None):
    sizes  = {1: 15, 2: 12.5, 3: 11}
    colors = {1: BLUE, 2: TEAL, 3: DARK}
    col = color or colors.get(level, BLACK)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = col
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '0084AA')
        pBdr.append(bottom); pPr.append(pBdr)
    return p

def bullet(text, level=0, size=11, color=BLACK, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix + "  "); r.bold = True
        r.font.size = Pt(size); r.font.color.rgb = TEAL
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color
    return p

def hline():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom); pPr.append(pBdr)

def ref_table(rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Inches(2.0), Inches(5.3)]
    hdr = tbl.rows[0]
    for cell, w, h in zip(hdr.cells, widths, ["Paper", "What to cite from it"]):
        cell.width = w; shade_cell(cell, '1A2E44')
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = WHITE
    for j, (paper, note) in enumerate(rows):
        row = tbl.rows[j + 1]
        bg = 'F4F6F8' if j % 2 == 0 else 'FFFFFF'
        for i, (cell, val, w) in enumerate(zip(row.cells, [paper, note], widths)):
            cell.width = w; shade_cell(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(val); run.font.size = Pt(9.5)
            if i == 0:
                run.bold = True; run.font.color.rgb = TEAL
            else:
                run.font.color.rgb = BLACK
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
run = p.add_run("Capstone Proposal — Work Division Plan")
run.bold = True; run.font.size = Pt(20); run.font.color.rgb = NAVY

para("Group 3  ·  DAMO 699  ·  Spring 2026",
     size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Tha Pyay Hmu  ·  Lhagii Tsogtbayar  ·  Nadia Ríos  ·  Jorge Mendoza  ·  Alrick Grandison",
     size=10.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
hline()

# ══════════════════════════════════════════════════════════════════════════════
# APA REMINDER BOX
# ══════════════════════════════════════════════════════════════════════════════
heading("\u26a0\ufe0f  Important: APA Style and Scientific Writing (Read Before Starting)")

reminders = [
    "Use APA 7th edition for all citations. In-text: (Author, Year) or Author (Year). "
    "Every claim that is not your own must be cited.",
    "All cited sources must appear in the References section at the end of the proposal "
    "in full APA format. Example: Ghodsi, A., Zaharia, M., et al. (2011). Dominant resource "
    "fairness. Proceedings of NSDI \u201811, 323\u2013336.",
    "Use scientific writing style: third person, objective tone, past tense for prior work "
    "('Doukha (2025) found that...'), present tense for your own claims ('our model predicts...').",
    "No bullet points in the final proposal paragraphs \u2014 write in complete sentences.",
    "Every table and figure must have a numbered caption.",
    "Refer to the Proposal Draft.md and Research_Paper_Reviews_Jobs.md files for paper "
    "details, numbers, and quotes to draw from.",
]
for rem in reminders:
    bullet(rem, size=10.5, color=DARK)

# ══════════════════════════════════════════════════════════════════════════════
# OVERVIEW TABLE
# ══════════════════════════════════════════════════════════════════════════════
heading("Section Assignments Overview")
para("Total content: \u22645 pages (excluding cover page and references).",
     size=10.5, color=GRAY, italic=True, space_after=4)
para("Note: 5 pages is tight for this technical project. Consider asking the instructor if "
     "6\u20137 pages is acceptable given the two heavy sections (5 and 6).",
     size=10, color=GOLD, italic=True, space_after=8)

overview_cols = ["Section", "Title", "Assigned To", "Est. Pages"]
overview_rows = [
    ("1", "Introduction", "Member 3", "~0.75"),
    ("2", "Problem Definition", "Member 3", "~0.50"),
    ("3", "Analytical Objective (Research Question + Hypothesis)", "Member 4", "~0.50"),
    ("4", "Proposed Data Sources", "Member 4", "~0.50"),
    ("5", "Proposed Analytical Approach", "Member 1 (Alrick)", "~1.50"),
    ("6", "Expected Outcomes and Contributions", "Member 2", "~0.75"),
    ("7", "Project Plan and Timeline", "Member 5", "~0.25"),
    ("8", "Ethical Considerations", "Member 5", "~0.25"),
    ("—", "References (APA)", "Member 5 assembles", "Not counted"),
    ("—", "Cover Page", "Member 5", "Not counted"),
    ("", "TOTAL", "", "~4.50\u20135.00"),
]

overview_widths = [Inches(0.6), Inches(3.4), Inches(1.9), Inches(1.4)]
tbl2 = doc.add_table(rows=1 + len(overview_rows), cols=4)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
hrow = tbl2.rows[0]
for cell, w, h in zip(hrow.cells, overview_widths, overview_cols):
    cell.width = w; shade_cell(cell, '0D1B2A')
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE

member_bg = ['E8F4FB', 'FDE8F4', 'E8F6EC', 'FFF3E0', 'F0E8FF',
             'E8F4FB', 'FDE8F4', 'E8F6EC', 'F0F0F0', 'F0F0F0', 'E0E0E0']
for j, (sec, title, member, pages) in enumerate(overview_rows):
    row = tbl2.rows[j + 1]
    bg = member_bg[j] if j < len(member_bg) else 'FFFFFF'
    for i, (cell, val, w) in enumerate(zip(row.cells, [sec, title, member, pages], overview_widths)):
        cell.width = w; shade_cell(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val); run.font.size = Pt(10)
        if i == 2 and "Alrick" in val:
            run.bold = True; run.font.color.rgb = MEMBER_COLORS[0]
        elif i == 2 and "Member" in val:
            idx = int(val.split()[1][0]) - 1
            run.font.color.rgb = MEMBER_COLORS[idx] if idx < len(MEMBER_COLORS) else DARK
        elif j == len(overview_rows) - 1:
            run.bold = True; run.font.color.rgb = NAVY
        else:
            run.font.color.rgb = BLACK

doc.add_paragraph()
hline()

# ══════════════════════════════════════════════════════════════════════════════
# MEMBER 1 — ALRICK
# ══════════════════════════════════════════════════════════════════════════════
heading("Member 1 — Alrick", level=1, color=MEMBER_COLORS[0])
heading("Section 5: Proposed Analytical Approach  (~1.5 pages)", level=2)

para(
    "This is the technical core of the proposal \u2014 the two-component analytical pipeline. "
    "Explain how the Predictive Model and the Optimization Model work together, justify the "
    "method choices, and describe how they would be implemented at capstone scope.",
    space_after=6
)

para("Required subsections (suggested):", bold=True, size=11, space_after=2)
bullet("Overview: the two-component pipeline (Predictive Model \u2192 feeds into \u2192 Optimization Model)")
bullet("Predictive Model: Random Forest, what it predicts (P95 peak, duration), features used, why RF over LSTM")
bullet("Optimization Model (Prescriptive): objective function, decision variables, constraints (memory, fairness, SLA deadline)")
bullet("Why this combination is appropriate: prediction enables safe overcommitment; optimization ensures no constraint is violated")
bullet("How it maps to Kubernetes: admission webhook, score phase, runtime controller")

para("Key papers to cite:", bold=True, size=11, space_before=8, space_after=4)
ref_table([
    ("Kovalenko & Zhdanova (2024)", "Mathematical model structure \u2014 binary assignment variables, capacity constraints, multi-objective form; the skeleton your model extends"),
    ("Resource Central \u2014 Cortez (2017)", "RF prediction architecture; tenant history as dominant feature; safe oversubscription result"),
    ("Doukha (2025)", "RF vs LSTM comparison (MAPE 2.65% vs 17.43%) \u2014 the head-to-head evidence justifying RF"),
    ("DRF \u2014 Ghodsi (2011)", "Fairness constraint formalization \u2014 dominant share definition and proven properties"),
    ("Coach \u2014 Reidys (2025)", "Temporal co-location strategy \u2014 additive pulse model, non-overlapping peaks"),
    ("Kofi (2025)", "Preprocessing pipeline (Savitzky-Golay + min-max); Google Cluster Trace v3 validation"),
    ("Priya (2025)", "Kubernetes scheduler plugin architecture \u2014 admission webhook + score + runtime controller"),
])

para("Writing guidance: The template says 'final model selection is not required \u2014 show analytical direction.' "
     "Present the approach as a justified plan. Use the objective function, feature set, and constraint list "
     "from Proposal Draft.md as the technical backbone. Write in third person: 'The proposed system trains a "
     "Random Forest model...'\u2019",
     size=10, color=GRAY, italic=True, space_after=4)

hline()

# ══════════════════════════════════════════════════════════════════════════════
# MEMBER 2
# ══════════════════════════════════════════════════════════════════════════════
heading("Member 2", level=1, color=MEMBER_COLORS[1])
heading("Section 6: Expected Outcomes and Contributions  (~0.75 page)", level=2)

para(
    "Describe what the project will produce, what improvements are expected, and what contribution "
    "this makes to the field. This section works from the outputs of Section 5 \u2014 you don\u2019t need to "
    "understand the model internals, just know: the system produces (a) a trained RF model, "
    "(b) a scheduling algorithm / simulation, and (c) results comparing against the default K8s scheduler.",
    space_after=6
)

para("Required subsections (suggested):", bold=True, size=11, space_after=2)
bullet("Expected quantitative outcomes with specific numbers (utilization %, SLA %, Gini, MAPE)")
bullet("What the system contributes that does not currently exist in the literature (the novelty claim)")
bullet("How practitioners benefit: cloud operators, tenants, organizations running shared clusters")
bullet("Limitations and scope: simulation only, single cluster, no live deployment")

para("Specific target numbers to include in the section:", bold=True, size=11, space_before=8, space_after=2)
bullet("Memory utilization: ~45\u201360% (default K8s) \u2192 \u2265 85% (target)", bold_prefix="\u2192 Target")
bullet("SLA violation rate: < 5%", bold_prefix="\u2192 Target")
bullet("Gini coefficient: ~0.25 \u2192 ~0.10", bold_prefix="\u2192 Target")
bullet("Prediction MAPE: < 5%", bold_prefix="\u2192 Target")

para("Key papers to cite:", bold=True, size=11, space_before=8, space_after=4)
ref_table([
    ("Priya (2025)", "Benchmark: <5% SLA violation rate \u2014 use as the bar to match or exceed"),
    ("Alatawi (2025)", "Gini coefficient improvement (0.25 \u2192 0.10) \u2014 your fairness target range"),
    ("Wang & Yang (2025)", "32.5% utilization improvement + 43.3% latency reduction \u2014 comparable result"),
    ("Heracles \u2014 Lo (2015)", "~90% utilization + <5% SLA is achievable \u2014 prior evidence targets are realistic"),
    ("Chaudhari (2025)", "The unbuilt framework \u2014 name your contribution explicitly as building what Chaudhari proposed"),
    ("Jiang Zhi (2025)", "Simulation on real traces is valid evaluation methodology \u2014 justify this approach"),
    ("Zhao et al. (2021)", "Profit = utilization framing \u2014 use for cloud provider stakeholder argument"),
])

hline()

# ══════════════════════════════════════════════════════════════════════════════
# MEMBER 3
# ══════════════════════════════════════════════════════════════════════════════
heading("Member 3", level=1, color=MEMBER_COLORS[2])
heading("Sections 1 + 2: Introduction + Problem Definition  (~1.25 pages total)", level=2)

para(
    "These two sections open the proposal. They must hook the reader and justify why the project "
    "matters. They flow together naturally \u2014 Introduction sets the industry context and stakes; "
    "Problem Definition precisely describes the analytical challenge.",
    space_after=6
)

para("Section 1 \u2014 Introduction (~0.75 page):", bold=True, size=11, space_after=2)
bullet("Context: multi-tenant cloud clusters, DRAM pricing pressure, AI workload growth")
bullet("Who the stakeholders are: cloud operators, tenants, organizations with shared clusters")
bullet("Why now: overcommitment is the right strategy; the scheduler is the bottleneck")

para("Section 2 \u2014 Problem Definition (~0.5 page):", bold=True, size=11, space_before=6, space_after=2)
bullet("The declared-vs-actual usage gap: tenants over-declare, creating wasted headroom")
bullet("Memory contention from simultaneous peaks: when multiple jobs peak at the same time, OOM kills happen")
bullet("Lack of native fairness: Kubernetes is first-come-first-served, no multi-tenant fairness")
bullet("How data analytics (prediction + optimization) is the right tool \u2014 not just engineering heuristics")

para("Key papers to cite:", bold=True, size=11, space_before=8, space_after=4)
ref_table([
    ("Chaudhari (2025)", "40\u201360% utilization waste in production; K8s fails for multi-tenant workloads \u2014 the gap anchor paper"),
    ("Borg \u2014 Verma (2015)", "Shared clusters save 20\u201330% infrastructure at scale; overcommitment is proven strategy"),
    ("Heracles \u2014 Lo (2015)", "Memory contention is the primary co-location interference; memory matters more than CPU"),
    ("Quasar \u2014 Delimitrou (2014)", "<20% CPU utilization, ~40% memory with static reservations \u2014 validates the problem"),
    ("DRF \u2014 Ghodsi (2011)", "Lack of multi-resource fairness as a structural gap in shared schedulers"),
    ("Pinnapareddy (2025)", "Kubernetes right-sizing gap; sustainability and cost motivation"),
])

para("Writing guidance: Use the Suggested Introduction paragraphs in Proposal Draft.md as your starting draft. "
     "Tighten and add APA citations. Every statistic ('%', target numbers) needs a citation.",
     size=10, color=GRAY, italic=True, space_after=4)

hline()

# ══════════════════════════════════════════════════════════════════════════════
# MEMBER 4
# ══════════════════════════════════════════════════════════════════════════════
heading("Member 4", level=1, color=MEMBER_COLORS[3])
heading("Sections 3 + 4: Analytical Objective + Proposed Data Sources  (~1.0 page total)", level=2)

para(
    "Section 3 states the research question and hypothesis. Section 4 describes the dataset. "
    "Together they answer: 'What are we trying to find out?' and 'What data will we use to find it?'",
    space_after=6
)

para("Section 3 \u2014 Analytical Objective (~0.5 page):", bold=True, size=11, space_after=2)
bullet("State the research question (see Option 1 in Proposal Draft.md \u2014 that is the working version)")
bullet("State the hypothesis with the four specific numeric targets")
bullet("Describe the type of project: predictive + prescriptive analytics combined")

para("Working research question to use:", bold=True, size=11, space_before=6, space_after=2)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.35); p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    '"By using the output of a predictive model as the input of an optimization model that is '
    'constrained by available resources and SLA requirements, we can optimally schedule workloads '
    'on a cluster such that resource idleness is kept minimal and fairness to tenants is high."'
)
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = DARK

para("Section 4 \u2014 Proposed Data Sources (~0.5 page):", bold=True, size=11, space_before=6, space_after=2)
bullet("Primary: Google Cluster Trace v3 \u2014 describe size, contents, public availability, why appropriate")
bullet("Optional secondary: Azure VM trace (Resource Central) \u2014 model generalization validation")
bullet("Justify why these sources match the analytical objective")

para("Key papers to cite:", bold=True, size=11, space_before=8, space_after=4)
ref_table([
    ("Kofi (2025)", "Google Cluster Trace v3 validated for this exact research type; preprocessing pipeline described"),
    ("Jiang Zhi (2025)", "Google Borg + Azure + Alibaba trace composition and size; multi-source trace comparison"),
    ("Resource Central \u2014 Cortez (2017)", "Azure VM trace; tenant history features available; similar RF training setup"),
    ("Wang & Yang (2025)", "Kubernetes multi-tenant scheduling as the deployment context for the dataset"),
    ("Doukha (2025)", "MAPE and MSE as the primary evaluation metrics for the prediction model"),
])

hline()

# ══════════════════════════════════════════════════════════════════════════════
# MEMBER 5
# ══════════════════════════════════════════════════════════════════════════════
heading("Member 5", level=1, color=MEMBER_COLORS[4])
heading("Sections 7 + 8: Timeline + Ethics + Final Integration", level=2)

para("Section 7 \u2014 Project Plan and Timeline (~0.25 page):", bold=True, size=11, space_after=2)
para("Describe the project phases using the 5-phase structure from the DAMO 699 template:", space_after=2)
bullet("Phase 1: Problem definition and proposal approval")
bullet("Phase 2: Data collection and preprocessing (Google Cluster Trace v3)")
bullet("Phase 3: Predictive model development (RF training + validation)")
bullet("Phase 4: Optimization model development (scheduling simulation + experiments)")
bullet("Phase 5: Results analysis, visualization, final report and presentation")
para("Include a rough timeline (weeks or months for each phase).", size=10.5, color=GRAY, italic=True, space_after=6)

para("Section 8 \u2014 Ethical Considerations (~0.25 page):", bold=True, size=11, space_before=6, space_after=2)
bullet("Confirm TCPS certification status for all group members")
bullet("Google Cluster Trace v3 is publicly available \u2014 no privacy concerns, no human subjects")
bullet("No proprietary data is used")
bullet("Results are for academic research only \u2014 not for production deployment without further validation")

para("Final Integration Responsibilities (not a section \u2014 coordination role):", bold=True, size=11, space_before=10, space_after=2)
bullet("Assemble all 5 members\u2019 sections into a single cohesive Word document")
bullet("Check for consistent terminology throughout: use 'Predictive Model' and 'Optimization Model' "
       "\u2014 avoid 'admission control' in the main narrative (it is a technical term; define it once if needed)")
bullet("Ensure section transitions flow and the document reads as one unified piece")
bullet("Build the References section in APA format by collecting all citations from all sections")
bullet("Check the cover page includes: title, course, program, names, group number, supervisor, institution, date")
bullet("Verify total page count \u2264 5 pages (content only, not cover or references)")

para("Papers likely cited across all sections (to include in References):", bold=True, size=11, space_before=8, space_after=4)
all_refs = [
    "Alatawi (2025) \u2014 serverless RL; Gini fairness metric",
    "Chaudhari (2025) \u2014 gap anchor paper; K8s fails for AI/ML",
    "Coach / Reidys (2025) \u2014 temporal co-location",
    "Delimitrou & Kozyrakis (2014) \u2014 Quasar; underutilization validation",
    "DRF / Ghodsi et al. (2011) \u2014 fairness algorithm",
    "Doukha (2025) \u2014 RF vs LSTM comparison",
    "Jiang Zhi (2025) \u2014 Clovers simulator; trace datasets",
    "Kovalenko & Zhdanova (2024) \u2014 math optimization model",
    "Kofi (2025) \u2014 Google Cluster Trace; preprocessing",
    "Liu & Guitart (2025) \u2014 in-node DRC enforcement",
    "Lo et al. (2015) \u2014 Heracles; memory contention evidence",
    "Perera (2025/2026) \u2014 RL model drift; landscape review",
    "Pinnapareddy (2025) \u2014 bin packing; Kubecost",
    "Priya (2025) \u2014 SloPolicy CRD; K8s scheduler plugin",
    "Resource Central / Cortez (2017) \u2014 RF prediction architecture",
    "Verma et al. (2015) \u2014 Borg; overcommitment at scale",
    "Wang & Yang (2025) \u2014 LSTM+DQN on Kubernetes",
    "Zhao et al. (2021) \u2014 admission control formalization",
]
for ref in all_refs:
    bullet(ref, size=10)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = r"C:\Users\alric\Documents\Spring2026\Capstone_ClaudeAI\Work_Division_Plan.docx"
doc.save(out)
print(f"Saved: {out}")
