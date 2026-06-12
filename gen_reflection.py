"""
gen_reflection.py
─────────────────
Generates the Capstone Reflection Paper (Week 11) as a Word document.

Format per assignment:
  Times New Roman, 12 pt, double-spaced, 1" margins, first person, headings.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent / "capstone_reflection_paper.docx"

doc = Document()

# ── Base formatting ─────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
pf = normal.paragraph_format
pf.line_spacing = 2.0
pf.space_after = Pt(0)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def heading(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def body(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def title_block() -> None:
    for line, bold in [
        ("Capstone Reflection Paper", True),
        ("Alrick Grandison", False),
        ("DAMO 699 Capstone Project", False),
        ("June 2026", False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2.0
        r = p.add_run(line)
        r.bold = bold
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)


# ── Content ─────────────────────────────────────────────────────────────────────
title_block()

# Brief project introduction (1-3 sentences) before the reflection begins.
_intro = doc.add_paragraph()
_intro.paragraph_format.line_spacing = 2.0
_intro.paragraph_format.first_line_indent = Inches(0.5)


def _irun(text: str, italic: bool = False) -> None:
    r = _intro.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = italic


_irun("My Capstone Project is titled ")
_irun("A Unified Framework for Multi-Tenant Cluster Management with Overcommit "
      "and Guaranteed Service Level Agreement", italic=True)
_irun(". It is a two-layer system that decides how to pack many tenants and their "
      "jobs onto a shared pool of machines while keeping memory use safe and meeting "
      "service level agreements. A prediction layer forecasts demand, and the "
      "optimization layer that I built turns those forecasts into real-time and "
      "plan-ahead scheduling decisions.")

heading("Initial Expectations and Goals")
body(
    "When I started the Capstone Project, I expected the optimization work to be easy. "
    "I thought a single, simple optimization model would solve the whole problem on its own. "
    "I pictured one clean formulation that I could write, run, and finish. I chose this "
    "Capstone because I want to work in the optimization analytics field after I finish school, "
    "and I felt this project would give me the hands-on experience I needed to get there."
)
body(
    "My personal goals were focused on skill building. I wanted to learn how to write "
    "optimization models for a complex system rather than for a textbook example. I wanted to "
    "understand how scientists and professionals actually develop optimization solutions for "
    "real world problems. Most of all, I wanted to learn the full process, which is how to "
    "reason about a problem, turn it into a model, build the model in code, and then test it."
)
body(
    "Looking back, my first expectation was the most important one to correct. The belief that "
    "one simple model would be enough set the tone for how I planned my early work. When that "
    "belief broke down, I had to learn to plan differently and to accept that a real solution "
    "would take several stages. That shift in mindset turned out to be one of the most valuable "
    "outcomes of the whole project."
)

heading("Role and Contributions")
body(
    "The Capstone Project had two layers. There was a prediction layer and an optimization "
    "layer. I chose to work on the optimization layer while the rest of the team worked on the "
    "prediction layer. My job was to create the mathematical models, write the code, and test "
    "the code that turned predicted demand into real scheduling decisions."
)
body(
    "My main responsibility was the Real-Time Optimization Model, which acts as the scheduler. "
    "I worked on the math for the model, developed the code, and tested it. I also took "
    "responsibility for refactoring the model when the analytical results showed it could be "
    "improved. Later I took on a second model, the Plan-Ahead Model, which we added to give the "
    "project more depth. I followed the same process for that model, starting from the math and "
    "ending with tested code. On top of the two models, I built a visual simulation that shows "
    "the entire pipeline running end to end, so that anyone could see how the prediction layer "
    "and the optimization layer fit together."
)

heading("Challenges and Problem-Solving")
body(
    "I faced challenges in communication and in technical limitations. Time management was not a "
    "major problem for me. The communication challenge was that my teammates sometimes did not "
    "understand the technical context I was explaining. The two layers used different ideas, and "
    "the optimization side carried a lot of math that was not obvious to everyone."
)
body(
    "The technical limitation appeared later, when I started doing deep computational time "
    "analysis on the models. I found that both models struggled when the parameters grew large. "
    "The solver took too long, and in some cases it ran out of memory before it finished."
)
body(
    "I addressed the communication challenge by sharing documents that broke down the complex "
    "topics into smaller pieces. When that was not enough, I met with my teammates and walked "
    "them through the ideas visually until they understood. I addressed the technical limitation "
    "by developing a decomposition heuristic in the code. Instead of solving one large problem, "
    "the heuristic breaks the problem into smaller pieces and solves them in sequence. This let "
    "the models run in a satisfactory time even when the parameters were large."
)
body(
    "Solving the limitation taught me more than solving the original model did. I had to study "
    "why the solver slowed down, test the models at different sizes, and measure where the cost "
    "grew the fastest. Only after I understood the cause could I design a fix that kept the "
    "results trustworthy while making the run time practical. That cycle of measure, diagnose, "
    "and redesign is a skill I will carry into any future optimization work."
)

heading("Learning and Development")
body(
    "I gained several new skills and tools during this project. I learned how to build a "
    "simulation dashboard using a backend and an API. I learned how to use commercial solvers "
    "such as Gurobi. I also gained an important insight, which is that real world problems are "
    "not solved by basic model math once the scale gets large. At scale, you need heuristics to "
    "get a good answer in a reasonable amount of time."
)
body(
    "The project greatly improved my understanding of the field. Multi-tenant memory "
    "utilization is a real world problem that businesses want to solve, because it controls "
    "cost and service quality at the same time. Working on it showed me how optimization "
    "analytics connects directly to business decision-making, and how a good model can turn a "
    "vague operational goal into a concrete and testable plan."
)

heading("Teamwork and Communication")
body(
    "I would evaluate our team dynamics as steady and reliable. My teammates were responsible "
    "for managing their part, which helped keep the project on pace. We communicated when we "
    "needed to, and we kept each other informed at the points where the two layers had to meet."
)
body(
    "I learned that teamwork can be challenging when you share ideas that not everyone "
    "understands. Having more people gets the work done faster, but it also creates more "
    "friction that the team has to work through. The lesson for me was that clear "
    "communication is just as important as technical skill, because a good model has no value "
    "if the rest of the team cannot follow how it works."
)

heading("Self-Assessment and Future Applications")
body(
    "I grew professionally during this project. I learned to take things one step at a time and "
    "to overcome problems incrementally instead of trying to solve everything at once. I started "
    "with one simple model, and I ended with two advanced models that use heuristics, plus a "
    "simulation dashboard that ties the whole pipeline together. That progression is the clearest "
    "sign of my growth."
)
body(
    "This experience will shape my future studies and my career. It has given me the confidence "
    "to put optimization expert on my resume as I look for a job after school. More than the "
    "title, it changed how I approach problems. I now expect real problems to be harder than "
    "they first appear, and I know how to break them down, build a solution, and test it until "
    "it works."
)

doc.save(str(OUT))
print(f"Reflection paper saved: {OUT}")
