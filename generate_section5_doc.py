from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(1.15); s.right_margin = Inches(1.15)

NAVY  = RGBColor(0x0D, 0x1B, 0x2A)
BLUE  = RGBColor(0x00, 0x6A, 0xA6)
TEAL  = RGBColor(0x00, 0x7A, 0x8A)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY  = RGBColor(0x60, 0x60, 0x60)

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14 if level == 1 else 12)
    r.font.color.rgb = BLUE if level == 1 else TEAL
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6'); bot.set(qn('w:color'), '006AA6')
        pBdr.append(bot); pPr.append(pBdr)

def para(text, bold=False, italic=False, size=11, color=BLACK, sb=0, sa=7):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size); r.font.color.rgb = color
    return p

# ── HEADER ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
r = p.add_run("Section 5: Proposed Analytical Approach")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

para("DAMO 699 Capstone Proposal  \u00b7  Group 3  \u00b7  Member 1 \u2014 Alrick Grandison",
     italic=True, size=10, color=GRAY, sa=12)

# ── OVERVIEW ─────────────────────────────────────────────────────────────────
para(
    "A multi-tenant application scheduler with a layered framework can provide fair distribution "
    "of workloads and high resource utilization (Chaudhari, 2025). This project proposes a "
    "framework with two layers: a prediction layer and an optimization layer. The prediction "
    "layer trains machine learning models on historical cluster data to forecast resource "
    "consumption by tenant workloads. The optimization layer takes these predictions as exogenous "
    "inputs and applies a formal optimization model to schedule workloads (Kovalenko & Zhdanova, "
    "2024). The framework simulates how a Kubernetes (K8s) scheduling component (Kubernetes, n.d.) "
    "would work but does not implement an actual K8s cluster, and serves as a base model for "
    "developing a production multi-tenant scheduler in future work. Memory is the target resource "
    "for increased utilization in this project; however, the framework includes both memory and "
    "Central Processing Unit (CPU) because CPU is required alongside memory for workload execution "
    "and is typically declared together with memory in container specifications. Network bandwidth, "
    "power consumption, and autoscaling are outside the scope of this project."
)

para(
    "The K8s default scheduler was designed for single-tenant cases (Priya, 2025). It assigns "
    "pods, containerized applications or workloads, to nodes, physical or virtual machines, "
    "based solely on declared resource requests. In a multi-tenant cluster, where multiple "
    "tenants (accounts or customers) share the same nodes, this approach causes systematic "
    "under-utilization and unequal resource access. This framework addresses both problems. "
    "It provides an orchestration model with a regularly updated prediction layer that feeds "
    "its outputs to an optimization layer. The optimization layer uses predictions, along with "
    "fairness and utilization objectives and resource constraints, to optimally schedule pods "
    "to nodes across tenants. The framework will train on open-source cluster data, specifically "
    "the Google Cluster Trace v3, and will run simulation on trace data as well as synthesized "
    "workload scenarios. This method works because we are learning the trends in how tenants "
    "deploy workloads and using that information to optimize resource utilization and scheduling "
    "decisions."
)

# ── PREDICTION LAYER ─────────────────────────────────────────────────────────
heading("Prediction layer", level=2)

para(
    "The prediction layer runs two models. The first model is a Long Short-Term Memory (LSTM) "
    "regressor for temporal predictions. Kofi (2025) showed that an LSTM model trained on the "
    "Google Cluster Trace v3 with Savitzky-Golay smoothing and min-max normalization preprocessing "
    "achieves an R\u00b2 of 0.99 for workload prediction, and significantly outperforms comparable "
    "models such as SVM and SATCN in prediction accuracy and error minimization. The LSTM captures "
    "the time-series behavior of tenant workloads and predicts when resource usage spikes will "
    "occur and how long they will last."
)

para(
    "The second model is a Random Forest (RF) classifier for job classification. Doukha and "
    "Ez-zahout (2025) compared RF against LSTM on CPU and memory prediction tasks. RF achieved a "
    "Mean Absolute Percentage Error (MAPE) of 2.65% while LSTM achieved 17.43% for static point "
    "predictions. RF is selected for classifying incoming jobs by resource consumption profile "
    "because it provides higher accuracy and interpretability than recurrent architectures for "
    "that task. Classification of job types also informs the optimization layer on which resource "
    "constraints apply to each workload class (Kovalenko & Zhdanova, 2024)."
)

para(
    "Tenants typically over-request memory as a precautionary buffer; actual workload consumption "
    "is consistently lower than declared. The prediction layer captures this gap. By using the "
    "95th percentile (P95) of predicted memory usage rather than the declared request, the "
    "framework can schedule more workloads onto a node than declared capacity would permit, safely "
    "overcommitting memory. This is a key contribution of the framework, as existing work tends to "
    "overcommit only CPU, where throttling is the worst-case outcome. Memory overcommitment carries "
    "higher risk: when a container exceeds its memory limit, the Linux kernel immediately terminates "
    "it in an Out-of-Memory (OOM) kill, ending the workload and constituting a Service Level "
    "Agreement (SLA) violation. CPU throttling only slows execution, but an OOM kill is "
    "unrecoverable. The P95 addresses this risk by covering 95% of real-world memory spikes without "
    "holding the full, over-declared reservation, providing a safety margin that makes memory "
    "overcommitment viable."
)

# ── OPTIMIZATION LAYER ───────────────────────────────────────────────────────
heading("Optimization layer", level=2)

para(
    "The optimization layer takes the P95 memory estimate and the temporal usage profile from "
    "the prediction layer as exogenous parameters (Kovalenko & Zhdanova, 2024). Rather than "
    "accepting tenant-declared resource requests as accurate, the optimization layer substitutes "
    "predicted actual consumption values. This substitution is the mechanism that enables safe "
    "overcommitment."
)

para(
    "Kovalenko and Zhdanova (2024) proposed a mathematical model for workload scheduling in "
    "cloud environments using a multi-objective function and resource capacity constraints. This "
    "project takes inspiration from that structure and extends it with prediction-driven inputs "
    "and a fairness objective. The objective function maximizes a weighted combination of three "
    "terms: average memory utilization across nodes, fairness equality across tenants, and SLA "
    "compliance rate. Resource capacity constraints ensure that the total predicted memory and "
    "CPU assigned to any node does not exceed that node\u2019s available capacity at any point in "
    "the simulation."
)

para(
    "The fairness objective draws from two sources. Beltre, Saha, and Govindaraju (2019) "
    "introduced Dominant Demand Share (DDS), which measures a tenant\u2019s resource share across "
    "both running workloads and its pending queue. Their experiments showed that DRF alone produced "
    "waiting time deviations of up to 157% above the cluster average, while combining DDS with DRF "
    "reduced this to below 13%. Shi and Yu (2024) formalize fairness as an optimization objective "
    "using Lyapunov optimization, where a virtual demand queue per tenant tracks unmet demand and "
    "its minimization ensures no tenant is chronically delayed, achieving 31.9% lower scheduling "
    "fairness variance than the best baseline. This project uses DDS as the fairness measure in "
    "the objective function and adapts the Job Scheduling Index (JSI) from Shi and Yu (2024), "
    "replacing their payment-bid demand signal with the RF-predicted memory consumption value."
)

para(
    "Together, the prediction layer and optimization layer form a framework that directly answers "
    "the research question. The prediction layer converts over-declared requests into accurate "
    "utilization estimates. The optimization layer then maximizes utilization and fairness "
    "simultaneously, subject to resource capacity constraints. Scheduling on predicted rather "
    "than declared values reduces idle capacity, and including fairness as an explicit objective "
    "ensures that utilization gains are distributed equitably across tenants.", sb=4
)

# ── REFERENCES ───────────────────────────────────────────────────────────────
heading("APA References for Section 5", level=2)

refs = [
    "Beltre, A., Saha, P., & Govindaraju, M. (2019). KubeSphere: An approach to multi-tenant "
    "fair scheduling for Kubernetes clusters. 2019 IEEE Cloud Summit, 14\u201320. "
    "https://doi.org/10.1109/CloudSummit47114.2019.00009",

    "Chaudhari, A. H. (2025). Multi-tenant AI workload scheduling on Kubernetes: Addressing "
    "modern cloud computing challenges. International Journal of Computational and Experimental "
    "Science and ENgineering, 11(3), 6870\u20136877. https://doi.org/10.22399/ijcesen.3931",

    "Doukha, R., & Ez-zahout, A. (2025). Enhanced virtual machine resource optimization in cloud "
    "computing using real-time monitoring and predictive modeling. International Journal of "
    "Advanced Computer Science and Applications, 16(2), 658. https://www.ijacsa.thesai.org",

    "Kofi, J. E. (2025). Data-driven cloud workload optimization using machine learning modeling "
    "for proactive resource management. International Journal of Emerging Research in Engineering "
    "and Technology, 6(4), 27\u201337. https://doi.org/10.63282/3050-922X.IJERET-V6I4P104",

    "Kovalenko, V., & Zhdanova, O. (2024). Dynamic mathematical model for resource management "
    "and scheduling in cloud computing environments. Information, Computing and Intelligent "
    "Systems, 5, 90\u2013100. https://doi.org/10.20535/2786-8729.5.2024/316545",

    "Kubernetes. (n.d.). kube-scheduler. Kubernetes Documentation. Retrieved April 19, 2026, "
    "from https://kubernetes.io/docs/concepts/scheduling-eviction/kube-scheduler/",

    "Priya, L. (2025). QoS-aware multi-tenant container orchestration using Kubernetes. "
    "International Journal of Advanced Research in Computer Science and Engineering, 1(3), 19\u201326. "
    "https://doi.org/10.63345/v1.i3.303",

    "Shi, Y., & Yu, H. (2024). Fairness-aware job scheduling for multi-job federated learning. "
    "arXiv:2401.02740v3. https://arxiv.org/abs/2401.02740",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(5)
    r = p.add_run(ref); r.font.size = Pt(10); r.font.color.rgb = BLACK

out = r"C:\Users\alric\Documents\Spring2026\Capstone_ClaudeAI\Section_5.docx"
doc.save(out)
print(f"Saved: {out}")
