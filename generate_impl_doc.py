from docx import Document
from docx.shared import Pt, RGBColor, Inches
import re

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

with open('suggested_implementation.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


def add_code_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text if text.strip() else ' ')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)


def add_rich_paragraph(doc, text, style_name='Normal'):
    p = doc.add_paragraph(style=style_name)
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text.strip())
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            p.add_run(part)
    return p


def render_table(doc, table_lines):
    rows = [r for r in table_lines if not re.match(r'^\s*\|[-| :]+\|\s*$', r)]
    if not rows:
        return
    ncols = rows[0].count('|') - 1
    if ncols < 1:
        return
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        for ci in range(ncols):
            cell_text = cells[ci] if ci < len(cells) else ''
            t.cell(ri, ci).text = cell_text
            if ri == 0:
                for para in t.cell(ri, ci).paragraphs:
                    for run in para.runs:
                        run.bold = True
    doc.add_paragraph()


in_code = False
code_buffer = []
table_buffer = []

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # Code block toggle
    if line.strip().startswith('```'):
        if not in_code:
            in_code = True
            code_buffer = []
        else:
            in_code = False
            for cl in code_buffer:
                add_code_line(doc, cl)
            doc.add_paragraph()
        i += 1
        continue

    if in_code:
        code_buffer.append(line)
        i += 1
        continue

    # Table rows
    if line.strip().startswith('|'):
        table_buffer.append(line)
        i += 1
        continue
    else:
        if table_buffer:
            render_table(doc, table_buffer)
            table_buffer = []

    stripped = line.strip()

    if stripped.startswith('#### '):
        add_heading(doc, stripped[5:], 4)
    elif stripped.startswith('### '):
        add_heading(doc, stripped[4:], 3)
    elif stripped.startswith('## '):
        add_heading(doc, stripped[3:], 2)
    elif stripped.startswith('# '):
        add_heading(doc, stripped[2:], 1)
    elif stripped == '---':
        p = doc.add_paragraph()
        p.add_run('_' * 55)
    elif stripped.startswith('- '):
        add_rich_paragraph(doc, stripped[2:], style_name='List Bullet')
    elif line.startswith('    ') and stripped:
        add_code_line(doc, line)
    elif stripped == '':
        doc.add_paragraph()
    else:
        add_rich_paragraph(doc, stripped)

    i += 1

# flush any remaining table
if table_buffer:
    render_table(doc, table_buffer)

doc.save('Suggested_Implementation.docx')
print('Done')
